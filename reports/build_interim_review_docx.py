"""Build the EEEM004 Interim Review Word document.

Outputs:
    reports/generated/exports/InterimReview.docx

The document follows the supervisor-facing Interim Review form structure,
with all student-owned sections populated from the rebuilt dissertation
framing. The blue supervisor boxes are intentionally left empty.

Run:
    venv/bin/python reports/build_interim_review_docx.py
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

import numpy as np
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
RESULTS = ROOT / "experiments" / "results"
EXPORTS = ROOT / "reports" / "generated" / "exports"
EXPORTS.mkdir(parents=True, exist_ok=True)


# --------------------------------------------------------------------------- #
# Helpers (kept parallel to build_main_dissertation_docx.py)
# --------------------------------------------------------------------------- #
def latest_json(prefix: str) -> dict | list:
    files = sorted(p for p in RESULTS.glob(f"{prefix}_*.json"))
    if not files:
        return []
    return json.loads(files[-1].read_text(encoding="utf-8"))


def avg(rows: Iterable[dict], key: str) -> float:
    rows = list(rows)
    if not rows:
        return float("nan")
    return float(np.mean([float(r[key]) for r in rows]))


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc: Document, text: str, *, italic: bool = False, bold: bool = False, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_supervisor_box(doc: Document, prompt: str) -> None:
    """Render a labelled empty box where the supervisor will write."""
    p = doc.add_paragraph()
    run = p.add_run(prompt)
    run.italic = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.size = Pt(10)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = " "
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EAF1F8")
    tc_pr.append(shd)
    set_cell_height(cell, Cm(2.5))


def set_cell_height(cell, height) -> None:
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height.emu / 635)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def set_default_font(doc: Document, family: str = "Calibri", size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = family
    style.font.size = Pt(size)
    rpr = style.element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), family)


def add_cover_table(doc: Document, fields: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(fields):
        lc, vc = table.rows[i].cells
        lc.text = label
        vc.text = value
        for run in lc.paragraphs[0].runs:
            run.bold = True
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_results_table(doc: Document, rows: list[dict]) -> None:
    headers = [
        "Agent",
        "Final value (USD)",
        "Sharpe",
        "Max DD",
        "VaR-95 viol.",
        "Terminal preservation",
        "Path preservation (1−MDD)",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for r in hdr_cells[i].paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["agent"]
        cells[1].text = f"${row['final']:,.0f}"
        cells[2].text = f"{row['sharpe']:+.4f}"
        cells[3].text = f"{row['mdd']:.4f}"
        cells[4].text = f"{row['var']:.4f}"
        cells[5].text = f"{row['pres']:.4f}"
        cells[6].text = f"{1.0 - float(row['mdd']):.4f}"
    for r in table.rows:
        for c in r.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_plan_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    headers = ["Working period", "Tasks to undertake", "Milestones (with target dates)"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for r in hdr_cells[i].paragraphs[0].runs:
            r.bold = True
    for period, tasks, milestones in rows:
        cells = table.add_row().cells
        cells[0].text = period
        cells[1].text = tasks
        cells[2].text = milestones
        for run in cells[0].paragraphs[0].runs:
            run.bold = True


def add_status_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    headers = ["Step", "Status", "Notes"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for r in hdr_cells[i].paragraphs[0].runs:
            r.bold = True
    for step, status, notes in rows:
        cells = table.add_row().cells
        cells[0].text = step
        cells[1].text = status
        cells[2].text = notes


# --------------------------------------------------------------------------- #
# Content assembly
# --------------------------------------------------------------------------- #
def build_results_rows() -> list[dict]:
    baseline = latest_json("baseline")
    prob = latest_json("probabilistic")
    bench = latest_json("benchmarks")
    rules = latest_json("rule_baseline")

    def _spy_only(rows: list) -> list:
        if not rows:
            return []
        return [r for r in rows if r.get("ticker", "SPY") == "SPY"]

    baseline = _spy_only(baseline)
    prob = _spy_only(prob)
    bench_spy = _spy_only(bench)
    rules_spy = _spy_only(rules)
    bench_lookup = {r["agent"]: r for r in bench_spy}
    rule_lookup = {r["agent"]: r for r in rules_spy}

    rows: list[dict] = []
    if baseline and prob:
        rows.extend([
            {
                "agent": "Baseline PPO",
                "final": avg(baseline, "final_portfolio_value"),
                "sharpe": avg(baseline, "sharpe_ratio"),
                "mdd": avg(baseline, "max_drawdown"),
                "var": avg(baseline, "var_95_violation_rate"),
                "pres": avg(baseline, "capital_preservation_rate_95pct_hwm"),
            },
            {
                "agent": "Probabilistic PPO",
                "final": avg(prob, "final_portfolio_value"),
                "sharpe": avg(prob, "sharpe_ratio"),
                "mdd": avg(prob, "max_drawdown"),
                "var": avg(prob, "var_95_violation_rate"),
                "pres": avg(prob, "capital_preservation_rate_95pct_hwm"),
            },
        ])
    for label, display in (
        ("stop_loss_5pct", "Rule-based stop-loss (5%)"),
        ("stop_loss_10pct", "Rule-based stop-loss (10%)"),
    ):
        r = rule_lookup.get(label)
        if r:
            rows.append({
                "agent": display,
                "final": float(r["final_portfolio_value"]),
                "sharpe": float(r["sharpe_ratio"]),
                "mdd": float(r["max_drawdown"]),
                "var": float(r["var_95_violation_rate"]),
                "pres": float(r["capital_preservation_rate_95pct_hwm"]),
            })
    for label in ("buy_and_hold", "all_cash"):
        b = bench_lookup.get(label)
        if b:
            rows.append({
                "agent": "Buy-and-hold (SPY)" if label == "buy_and_hold" else "All-cash",
                "final": float(b["final_portfolio_value"]),
                "sharpe": float(b["sharpe_ratio"]),
                "mdd": float(b["max_drawdown"]),
                "var": float(b["var_95_violation_rate"]),
                "pres": float(b["capital_preservation_rate_95pct_hwm"]),
            })
    return rows


# --------------------------------------------------------------------------- #
# Build
# --------------------------------------------------------------------------- #
def build() -> Path:
    doc = Document()
    set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ----- Cover banner -----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("EEEM004 — MSc Project")
    r.bold = True
    r.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Interim Review")
    r.italic = True
    r.font.size = Pt(14)

    doc.add_paragraph()
    add_para(
        doc,
        "This document is the student's submission for the EEEM004 interim review. "
        "It mirrors the structure of the official Interim Review form. The blue boxes "
        "are reserved for the supervisor's written assessment and are intentionally "
        "left empty.",
        italic=True,
    )

    # ----- Cover information -----
    add_heading(doc, "Cover information", 1)
    add_cover_table(doc, [
        ("Name", "Fiyin Akano"),
        ("URN", "6962514"),
        ("Supervisor", "Dr Cuong Nguyen"),
        ("Second supervisor (if applicable)", "[INSERT IF APPLICABLE]"),
        ("Date of meeting", "[INSERT MEETING DATE]"),
        ("Module", "EEEM004 — MSc Dissertation (cross-year)"),
        ("Department", "Electrical and Electronic Engineering, University of Surrey"),
        ("Target submission date", "1 September 2026"),
    ])

    # ----- Project title -----
    add_heading(doc, "Project title", 1)
    add_para(
        doc,
        "Probabilistic Deep Reinforcement Learning for Portfolio Risk Analysis: "
        "drawdown-constrained portfolio control with an uncertainty-aware "
        "reinforcement-learning policy.",
        bold=True,
    )

    # ----- Problem statement -----
    add_heading(doc, "Problem statement", 1)
    add_para(
        doc,
        "Start with a concrete picture. Imagine putting one million US dollars into "
        "the US stock market in January 2022 and holding on. By January 2025 the "
        "account is worth about $1.52 million — a clean win. But on the way there, "
        "in October 2022, the same account briefly read $750,000 — a 25 % drop from "
        "the peak. For an individual that drop is scary; for a pension fund or an "
        "endowment it is something different — it is a breach of contract. Many "
        "institutional mandates carry an explicit drawdown limit (a maximum permitted "
        "loss measured from the peak rather than from the starting balance), and "
        "when that limit is breached, redemption rights kick in, trustees can be "
        "removed, and the fund can be forcibly liquidated. Buy-and-hold violates "
        "these limits routinely; manually setting a stop-loss (sell when the price "
        "has fallen 5 % below its peak) violates them too — it sells late and buys "
        "back later. The dissertation asks whether a small AI agent can do better.",
    )
    add_para(
        doc,
        "Reframed problem statement. Many investors and institutional mandates are "
        "required to keep portfolio drawdown from peak below a stated limit (commonly "
        "between 5 % and 20 %) while still beating cash and ideally beating passive "
        "index exposure. The standard quantitative answers — Markowitz mean-variance, "
        "risk parity, fixed-rule stop-losses — either assume the joint distribution "
        "of returns is stationary or react too slowly when the market regime changes. "
        "This dissertation studies whether a deep-reinforcement-learning agent that "
        "conditions on its own forecaster's predictive uncertainty (how confident the "
        "forecaster is, not just what it predicts) can sit on a more attractive point "
        "of the return-versus-drawdown trade-off than (a) passive buy-and-hold, (b) "
        "a rule-based stop-loss policy, and (c) a baseline PPO with no uncertainty "
        "signal — measured on a held-out test window that contains real macro shocks "
        "(2022–2025), with reproducible random seeds and an out-of-time generalisation "
        "check.",
    )
    add_para(
        doc,
        "Where the standard answers fail. Mean-variance (Markowitz, 1952) is single-"
        "period and treats upside and downside wiggles symmetrically — it has no "
        "memory of the running peak. Value-at-Risk and expected shortfall "
        "(Rockafellar and Uryasev, 2000) measure how bad a single bad day could be, "
        "but they cannot tell you how many bad days in a row you can endure before "
        "the loss-from-peak crosses the limit. Conditional drawdown-at-risk "
        "(Chekhlov, Uryasev and Zabarankin, 2005) is path-dependent and matches the "
        "institutional constraint shape, but it is a one-shot static optimisation: "
        "it picks one weight vector and does not adapt mid-window. Reactive trailing "
        "stop-losses adapt sequentially but fire after the drawdown has already "
        "begun and typically forfeit the recovery on the way back up.",
    )
    add_para(
        doc,
        "Why this matters in practice. Drawdown control is a real, billion-dollar "
        "institutional problem. CalPERS and other major pension funds, sovereign-"
        "wealth funds, university endowments and CTA funds all run explicit drawdown "
        "limits in their governance documents. Bridgewater Associates' All Weather "
        "fund, which has managed over 150 billion US dollars at peak, is publicly "
        "described by its founder as designed to lose less in any environment — an "
        "explicit drawdown-control objective. The drawdown literature itself "
        "(Magdon-Ismail and Atiya, 2004; Young, 1991; Sortino and Price, 1994; "
        "Chekhlov, Uryasev and Zabarankin, 2005) is the formal home for these "
        "constraints; this dissertation does not invent them, it picks up the "
        "tradition and extends it from one-shot optimisation to a sequential, "
        "uncertainty-aware decision policy.",
    )

    # ----- Objectives -----
    add_heading(doc, "Objectives", 1)
    add_para(
        doc,
        "The objectives have been refined during Phase 0 and Phase 1 in light of "
        "supervisor feedback. The current working set is:",
    )
    add_bullets(doc, [
        "O1. To study whether an explicit forecast-uncertainty signal, modelled with a DeepAR-style probabilistic LSTM and consumed by a PPO policy as both a state feature and a hard guard on new long-side actions, allows the agent to sit closer to the return-versus-drawdown frontier than uncertainty-blind alternatives on US equity index data.",
        "O2. To evaluate the resulting policy on a held-out window containing real macro shocks (2022 to 2025) against three named comparators — passive buy-and-hold, a rule-based stop-loss policy, and a baseline PPO with no uncertainty signal — using a metric set in which the headline criteria are Sharpe ratio, terminal value relative to buy-and-hold, and the capital-preservation ratio against the running high-watermark.",
        "O3. To pin down a fully reproducible evaluation protocol of fixed splits, fixed seeds, scripted artefacts and a shared metric set, so that any comparison made in this dissertation is genuinely like-for-like and can be reproduced from the public repository in a single command sequence.",
        "O4. To take an honest position, on the strength of O1 to O3, on when an uncertainty signal earns a place in a portfolio control loop and, just as important, on when it does not.",
    ])

    page_break(doc)

    # ----- Literature -----
    add_heading(doc, "Literature review (key references)", 1)
    add_para(
        doc,
        "The annotated list below is the working set of references that anchor the "
        "dissertation. Each entry follows a uniform structure — what the paper did, "
        "why they did it, what it does well, where it falls short, and how the gap is "
        "fixed in the literature — so the comparative differentiation Dr Nguyen asked "
        "for is visible at a glance. The dissertation's full bibliography in Chapter 2 "
        "of Main_Dissertation_Draft.docx is the canonical version; the list below is "
        "the compact working set for this interim review.",
    )
    add_para(
        doc,
        "Why the families differ at a glance. The risk-measure literature is not a "
        "ranking of competing answers to one question; it is a sequence of answers to "
        "different questions, each developed in response to a specific limitation of "
        "the family that came before:",
    )
    add_bullets(doc, [
        "Markowitz (1952) — \"What is the smoothest portfolio I can build given my return target?\" Single-period; symmetric penalty on wiggles.",
        "Rockafellar and Uryasev (2000) — \"What is the portfolio with the smallest expected loss in the worst 5 % of single-period outcomes?\" Single-period; coherent tail measure. Differs from Markowitz in shifting from average wiggle to tail loss.",
        "Magdon-Ismail and Atiya (2004) / Young (1991) / Chekhlov, Uryasev and Zabarankin (2005) — \"What is the worst loss-from-peak this strategy can be expected to experience over its lifetime?\" Path-dependent. Differs from VaR/ES in shifting from a single bad day to the worst peak-to-trough excursion of the equity curve.",
        "Sortino and Price (1994) — \"How much per-period downside deviation does this strategy accept to earn its return?\" Same per-period grain as Markowitz but with the symmetric penalty replaced by an asymmetric one. Differs from Markowitz in its treatment of upside.",
        "Jiang/Yang/Liu DRL-finance papers — \"Can a deep RL policy outperform static rules on equities and crypto with return-only reward?\" Sequential decision but uncertainty-blind. Differs from the risk-measure tradition in treating risk as something measured after the fact rather than constrained explicitly.",
    ])
    add_para(
        doc,
        "Drawdown literature (the closest antecedents).",
        bold=True,
    )
    refs = [
        ("Magdon-Ismail and Atiya (2004) — \"Maximum Drawdown.\" Risk Magazine, 17(10), 99–102.",
         "What and why: derived a closed-form expectation for the maximum drawdown of a geometric Brownian motion with drift, motivated by practitioner demand for an analytical baseline against which observed MDD figures could be calibrated. Advantages: closed-form, operationally meaningful, tied directly to (mu, sigma, T). Disadvantages: requires Brownian motion with constant drift and volatility — empirically wrong for financial returns, which exhibit volatility clustering and fat tails. Fix: extends to GARCH and regime-switching baselines, or — operationally — to the CDaR programme of Chekhlov, Uryasev and Zabarankin (2005). Position in this dissertation: defines the maximum drawdown statistic used throughout Section 5–Section 6 and motivates the analytical context for the test-window MDD numbers."),
        ("Young (1991) — \"Calmar Ratio: A Smoother Tool.\" Futures, 20(1), 40.",
         "What and why: introduced the Calmar ratio (annualised return / maximum drawdown) as a return-per-unit-pain measure intuitive to institutional investors with explicit drawdown limits. Advantages: no parametric assumption on returns; speaks the language of the institutional mandate; widely understood. Disadvantages: extreme single-event sensitivity — one bad week can permanently impair a strategy's Calmar; rewards strategies whose worst loss is still ahead of them. Fix: Sterling and Burke ratios use averages of the worst N drawdowns; the CDaR programme of Chekhlov, Uryasev and Zabarankin (2005) uses a tail expectation over the full drawdown distribution and removes the single-event problem. Position in this dissertation: reported as a diagnostic alongside Sharpe in Section 5.5; its limitations are exactly what motivated the move to the constrained CDaR-style objective in Section 3.1.5."),
        ("Chekhlov, Uryasev and Zabarankin (2005) — \"Drawdown Measure in Portfolio Optimization.\" International Journal of Theoretical and Applied Finance, 8(1), 13–58.",
         "What and why: extended the CVaR construction of Rockafellar and Uryasev (2000) from the loss distribution to the path-dependent drawdown distribution, defining the conditional drawdown-at-risk (CDaR_alpha) and showing that drawdown-constrained portfolio optimisation reduces to a linear programme under empirical scenarios. Motivated by institutional allocators (endowments, sovereign wealth funds) who already demanded CDaR diagnostics but had no tractable optimisation routine for them. Advantages: coherent generalisation of MDD/Calmar to a tail expectation over the drawdown distribution; LP-tractable; matches institutional mandates. Disadvantages: computational cost grows with horizon T; remains a one-shot static optimisation that does not naturally accommodate a sequential decision policy. Fix: multi-stage stochastic programming (Bertsimas, Lauprete and Samarov, 2004) extends to multiple periods; reinforcement learning extends further to a state-conditional policy that learns when to reduce exposure rather than only how to allocate it. Position in this dissertation: the closest classical antecedent to the constrained problem in Section 3.1.5. The dissertation's contribution can be stated as enforcing the CDaR-style constraint inside a sequential PPO policy rather than at portfolio-construction time."),
        ("Sortino and Price (1994) — \"Performance Measurement in a Downside Risk Framework.\" Journal of Investing, 3(3), 59–64.",
         "What and why: replaced the standard deviation in the Sharpe-ratio denominator with the downside deviation against a target return tau, motivated by Sortino's long-running argument that Sharpe penalises long-only mandates twice over (once for negative returns, once for above-average positive returns). Advantages: asymmetric penalty matches asymmetric utility of long-only investors; same data requirement as Sharpe; widely understood in CTA and hedge-fund evaluation. Disadvantages: tau is a modelling decision rather than a derived quantity, so the same strategy can have very different Sortino ratios under different rate environments; unstable when there are few sub-tau returns; remains single-period and silent on the path of the equity curve. Fix: report at multiple tau; pair with a path-dependent measure such as MDD or Calmar; lower-partial-moment family (Bawa, 1975; Fishburn, 1977) generalises the squared deviation. Position in this dissertation: reported in the Section 3.7 metric table; the asymmetric-penalty intuition behind Sortino is the same intuition behind the choice of drawdown rather than variance as the binding measure."),
    ]
    for r in refs:
        add_para(doc, r[0], bold=True)
        add_para(doc, r[1])
    add_para(
        doc,
        "Classical risk-measure literature (single-period baselines).",
        bold=True,
    )
    refs = [
        ("Markowitz (1952) — \"Portfolio Selection.\" Journal of Finance, 7(1), 77–91.",
         "What and why: turned portfolio choice into a quadratic optimisation in mean and variance, motivated by the absence of any quantitative framework for diversification before 1952. Advantages: convex QP solvable in milliseconds; closed-form efficient frontier in the unconstrained case; foundation of every later extension. Disadvantages: requires stationary mu and Sigma; weights are highly sensitive to estimation error (Michaud, 1989, \"error-maximising\"); penalises upside and downside symmetrically; single-period and ignores path properties. Fix: Black-Litterman shrinkage; Ledoit-Wolf covariance shrinkage; replace symmetric variance with downside deviation (Sortino) or tail loss (CVaR) or drawdown (CDaR). Position here: the classical static baseline that the sequential, uncertainty-aware policy is positioned against."),
        ("Rockafellar and Uryasev (2000) — \"Optimization of Conditional Value-at-Risk.\" Journal of Risk, 2(3), 21–42.",
         "What and why: wrote down a convex programme whose minimiser is the CVaR of the portfolio's loss distribution, motivated by the post-Artzner et al. (1999) recognition that VaR is not coherent and the practitioner need for a tractable, optimisable alternative. Advantages: coherent (sub-additive); convex in the weights; reduces to an LP under empirical scenarios. Disadvantages: single-period, so silent on the path of the equity curve; tail-data-hungry. Fix: extreme-value-theory parametric fits for the tail (McNeil and Frey, 2000); path-dependent generalisation in CDaR (Chekhlov et al., 2005). Position here: motivates the tail-loss arm of Section 2.1 and the VaR-95 violation-rate metric reported in Section 3.7."),
    ]
    for r in refs:
        add_para(doc, r[0], bold=True)
        add_para(doc, r[1])

    add_para(
        doc,
        "Reinforcement learning, policy gradient and DRL-for-finance.",
        bold=True,
    )
    refs = [
        ("Sutton and Barto (2018) — Reinforcement Learning: an Introduction, 2nd ed., MIT Press.",
         "What and why: the canonical textbook treatment of MDPs and policy-gradient methods, motivated by twenty years of fragmented research that needed a single integrated reference. Advantages: standardises notation and conceptual scaffolding for the entire field; covers tabular and function-approximation methods uniformly. Disadvantages: predates the deep-RL revolution in its proportions, so coverage of large-scale function approximation is comparatively light. Fix: pair with Schulman et al. (2017) for modern policy-gradient practice. Position here: the reference relied on for the MDP formulation and the policy-gradient derivations in Section 3.5."),
        ("Schulman, Wolski, Dhariwal, Radford and Klimov (2017) — \"Proximal Policy Optimization Algorithms.\" arXiv:1707.06347.",
         "What and why: introduced a clipped surrogate objective that approximates a trust-region constraint without the conjugate-gradient machinery of TRPO, motivated by the need for an algorithm with TRPO-like performance and SGD-like operational simplicity. Advantages: forgiving with respect to hyper-parameters; first-order optimisation only; mature open-source implementations (Stable-Baselines3). Disadvantages: clipping is a heuristic with no formal trust-region guarantee; sensitive to advantage-estimation method; ratio variable can be numerically unstable on long horizons. Fix: PPO-EWMA (Hilton et al., 2022); trust-region-aware clipping. Position here: PPO is the policy-gradient algorithm used for both the baseline and the uncertainty-aware variant. The clipped surrogate is the reason PPO was picked over vanilla policy gradient or TRPO: the operational simplicity matters on a single-asset environment with limited training budget."),
        ("Jiang, Xu and Liang (2017) — \"A Deep Reinforcement Learning Framework for the Financial Portfolio Management Problem.\" arXiv:1706.10059.",
         "What and why: treated portfolio management as an end-to-end DRL problem on cryptocurrency, motivated by the inadequacy of static rules on a non-stationary high-frequency market. Advantages: genuinely end-to-end (raw price tensor → portfolio weights); validated on a market where mean-variance baselines fail; established the template later adopted by FinRL. Disadvantages: return-only reward; cryptocurrency-specific; uncertainty-blind; tested on a single window without walk-forward. Fix: Yang et al. (2020) and FinRL extend to equities and add reproducibility infrastructure; this dissertation closes the uncertainty-blind and risk-blind gaps. Position here: the closest prior DRL-finance work in spirit, against which the present dissertation's contribution is positioned in Section 2.4."),
        ("Yang, Liu, Zhong and Walid (2020) — \"Deep Reinforcement Learning for Automated Stock Trading: An Ensemble Strategy.\" ICAIF.",
         "What and why: applied the Jiang template to US equities and added an algorithm-level ensemble that picks the best of A2C/PPO/DDPG per regime, motivated by the high run-to-run variance of any single policy-gradient algorithm. Advantages: equities-validated; engineered technical indicators give useful inductive bias; explicit train/validation/test split protocol that this dissertation follows. Disadvantages: still return-only; technical-indicator features embed prior beliefs; in-sample regime detection has look-ahead concerns; uncertainty-blind. Fix: walk-forward evaluation (Bailey and López de Prado, 2014); the present dissertation's uncertainty coupling is orthogonal to the algorithm-ensemble idea and can in principle be added on top. Position here: borrowed the train/validation/test split style and the seed-averaging convention; positioned as an algorithm-level rather than constraint-level contribution."),
        ("Liu, Yang, Gao and Wang (2021) — FinRL. arXiv:2011.09607.",
         "What and why: open-source library bundling the trading environment, data pipeline, algorithm interface and reporting layer, motivated by the divergent and incomparable results in the DRL-finance literature up to 2020. Advantages: imposes a single environment specification on the field; widely adopted; gym-style API. Disadvantages: takes no position on what the policy should optimise; bundles implicit environment choices that are not the right fit for every problem; abstraction layer hides design decisions that need to be defended at viva. Fix: implement the environment explicitly when the contribution is on the environment side rather than the policy side. Position here: this dissertation does not subclass FinRL; it implements its own environment in experiments/common.py so the trade-scaling and entry-guard contributions are explicit."),
        ("Raffin, Hill, Gleave, Kanervisto, Ernestus and Dormann (2021) — \"Stable-Baselines3.\" JMLR, 22(268), 1–8.",
         "What and why: a maintained, tested PyTorch port of the Stable-Baselines RL library, motivated by reproducibility and quality concerns with the original TF1 implementations. Advantages: rigorously tested; conforms to the gym/gymnasium API; interoperable with the wider Python RL ecosystem. Disadvantages: opinionated defaults that occasionally need to be overridden for non-standard environments; not the right tool for novel-algorithm research where every line of the policy needs to be inspected. Fix: subclass the algorithm class to override default hyper-parameters; for novel-algorithm work, write the algorithm from scratch and keep SB3 as the reproducibility baseline. Position here: the implementation source for the PPO solver in both runners."),
    ]
    for r in refs:
        add_para(doc, r[0], bold=True)
        add_para(doc, r[1])

    add_para(
        doc,
        "Probabilistic forecasting and uncertainty quantification.",
        bold=True,
    )
    refs = [
        ("Salinas, Flunkert, Gasthaus and Januschowski (2020) — \"DeepAR: Probabilistic Forecasting with Autoregressive Recurrent Networks.\" International Journal of Forecasting, 36(3), 1181–1191.",
         "What and why: an autoregressive RNN trained to maximise the likelihood of observed data under a chosen output distribution, motivated by Amazon's production demand-forecasting need for full predictive distributions and statistical pooling across thousands of related series. Advantages: emits a calibrated predictive distribution as a native output; pools strength across related series; trains with standard backprop. Disadvantages: parametric output family (Gaussian) imposes thin-tailed conditional residuals that financial returns violate; can overfit short series; calibration on out-of-sample regimes is not guaranteed; provides aleatoric but not epistemic uncertainty. Fix: Student-t or mixture heads; deep ensembles for the epistemic component; conformal prediction for non-parametric calibration. Position here: the blueprint for the forecaster in this dissertation. The implementation is a stripped-down DeepAR with a Gaussian head emitting predictive mean and log-variance, used to feed a one-dimensional uncertainty signal u_t into the PPO policy."),
        ("Hochreiter and Schmidhuber (1997) — \"Long Short-Term Memory.\" Neural Computation, 9(8), 1735–1780.",
         "What and why: introduced the LSTM cell with input/forget/output gates and a constant-error-carousel state, motivated by the inability of vanilla RNNs to learn long-range dependencies because of vanishing gradients. Advantages: the architectural workhorse for sequence modelling for two decades; trains stably with standard backpropagation through time; well understood. Disadvantages: superseded on many tasks by Transformer-based architectures (Vaswani et al., 2017) which scale better with data and compute; recurrent computation does not parallelise across time. Fix: Transformer-based forecasters (Informer, Autoformer); for short series the LSTM is still competitive. Position here: the architectural workhorse inside the forecaster."),
        ("Lakshminarayanan, Pritzel and Blundell (2017) — \"Simple and Scalable Predictive Uncertainty Estimation using Deep Ensembles.\" NeurIPS.",
         "What and why: trained several independent networks with different random initialisations and used their disagreement on a held-out point as the uncertainty proxy, motivated by the operational complexity of full Bayesian neural networks. Advantages: implementation-agnostic; captures both data-side and model-side uncertainty; well calibrated empirically. Disadvantages: training cost scales linearly with ensemble size; the disagreement signal is unstructured; storage cost. Fix: parameter-efficient ensembles (BatchEnsemble, MIMO); single-model alternatives such as MC dropout or evidential heads. Position here: read to decide whether to go with ensembles or a single network with a Gaussian head. The single-head approach is cheaper and sufficient given the policy only needs a one-dimensional uncertainty summary; ensembles are listed as future work."),
        ("Gal and Ghahramani (2016) — \"Dropout as a Bayesian Approximation.\" ICML.",
         "What and why: showed that dropout at inference time can be interpreted as variational approximation to a Bayesian posterior over weights, motivated by the wish to extract uncertainty estimates from existing deep models without architectural change. Advantages: zero training-time overhead; works on existing trained networks; gives empirical predictive distributions cheaply. Disadvantages: variational approximation is loose; underestimates epistemic uncertainty in some settings; calibration depends on dropout rate. Fix: combine with last-layer Bayesian methods or with conformal calibration. Position here: cited as background reading for why a learned predictive variance can be treated as a meaningful uncertainty proxy; the dissertation does not use MC dropout because the DeepAR-style head provides the same signal more cheaply at training time."),
        ("Vovk, Gammerman and Shafer (2005) — Algorithmic Learning in a Random World. Springer.",
         "What and why: developed conformal prediction, a non-parametric, distribution-free framework for constructing finite-sample-valid prediction sets, motivated by the desire for calibration guarantees that do not depend on parametric distributional assumptions. Advantages: model-agnostic; finite-sample validity under exchangeability; can wrap any underlying predictor. Disadvantages: predictive intervals can be conservative; standard exchangeability assumption is violated by time-series data. Fix: split-conformal and adaptive variants for non-exchangeable data (Gibbs and Candès, 2021). Position here: noted in Section 2.6 and Section 7.2 as a future-work overlay that would give the dissertation's uncertainty signal a finite-sample coverage guarantee independent of the Gaussian-head assumption."),
    ]
    for r in refs:
        add_para(doc, r[0], bold=True)
        add_para(doc, r[1])

    page_break(doc)

    # ----- Technical progress -----
    add_heading(doc, "Technical progress", 1)
    add_heading(doc, "Summary", 2)
    add_para(
        doc,
        "Phase 0 and Phase 1 are in place and reproducible end-to-end. The "
        "dissertation compares a baseline PPO agent against a probabilistic-PPO "
        "variant that consumes a DeepAR-style uncertainty signal, with passive "
        "buy-and-hold, all-cash, and a rule-based trailing stop-loss policy as the "
        "three named comparators. Everything is driven from a single configuration "
        "file and a small set of scripts, so any reported number can be regenerated "
        "from the public repository in a single command sequence.",
    )

    add_heading(doc, "What has been built", 2)
    add_bullets(doc, [
        "A probabilistic forecaster (experiments/run_probabilistic_agent.py): an LSTM trained with Gaussian negative log-likelihood that emits the mean and log variance of the next-step log return. The predictive standard deviation is min-max normalised across the test window into a unit-interval uncertainty score.",
        "An uncertainty-aware trading environment (experiments/common.py:StockEnv): action space [-1, 1] over a configurable max_trade_fraction of cash, with the trade size shrunk by (1 - uncertainty_level) and floored at min_trade_scale. When the uncertainty score exceeds the protocol quantile (default 0.80) the environment blocks new long-side trades but still allows exits. Reward is the per-step log of the portfolio-value ratio multiplied by 100 for numerical scale.",
        "A baseline PPO runner (experiments/run_baseline.py) that uses the same environment without the uncertainty coordinate or the trade-size shrinkage, so the comparison against the probabilistic variant is genuinely controlled.",
        "A rule-based stop-loss runner (experiments/run_rule_baselines.py) implementing 5 % and 10 % trailing-stop policies with a 20/50-day moving-average crossover for re-entry. This is the directly-measured non-AI comparator that the supervisor's previous feedback asked for.",
        "A benchmarks runner (experiments/run_benchmarks.py) that evaluates passive buy-and-hold and all-cash on the same test window. These act as sanity checks on the metric definitions as much as competitors to beat.",
        "A single evaluation protocol (experiments/configs/dissertation_protocol.json) that fixes the splits (2009-2018 train / 2019-2021 validation / 2022-2025 test), the seeds [7, 19, 42] in the headline study, and the metric set, and is read by every script. This is the bit that actually makes the comparisons fair.",
        "A reporting layer: reports/generate_dissertation_report.py for the markdown summary, reports/build_supervisor_pack.py for the one-page chart, reports/plot_dissertation_visuals.py for the detailed figures, and Dissertation_Walkthrough.ipynb for the embedded-output review notebook.",
    ])

    add_heading(doc, "Phase-0 to Phase-1 status table", 2)
    add_status_table(doc, [
        ("0.1 Environment + dependencies", "Done", "requirements.txt, SB3, PyTorch, gymnasium, yfinance"),
        ("0.2 PPO baseline on sample data", "Done", "phase0_examples/ppo_stock_trading_standalone.py"),
        ("0.3 DeepAR-style probabilistic example", "Done", "phase0_examples/deepar_style_example.py"),
        ("1.1 Shared protocol + metrics", "Done", "experiments/configs/dissertation_protocol.json, experiments/common.py"),
        ("1.2 Reproducible baseline / probabilistic / benchmark runners", "Done", "Three runners, seeded"),
        ("1.3 Dissertation report + supervisor pack", "Done", "reports/generated/"),
        ("1.4 Rule-based stop-loss comparator (5 % and 10 % variants)", "Done", "experiments/run_rule_baselines.py"),
        ("1.5 Robustness on 70-ticker test universe (Phase-1 budget)", "Done", "Four-agent comparison on 70 tickers × 3 seeds × 10k steps; aggregate stats + per-ticker table in Section 5.5 of dissertation"),
        ("1.6 Walk-forward (out-of-time) on CPU-feasible subset", "Done", "4 tickers × 4 folds × 3 seeds × 10k steps = 96 trainings; in Section 6.4 of dissertation"),
        ("1.7 Extended seed-stability check on representative sub-universe", "Done", "8 tickers × 10 seeds × 50k steps = 80 trainings; in Section 5.5.1 of dissertation"),
        ("1.8 Phase-2 extended grid on full 70-ticker universe", "Scheduled", "GPU-only; orchestrator experiments/run_extended_grid.py + notebook notebooks/extended_grid_colab.ipynb"),
    ])

    add_heading(doc, "Current results (mean across 3 seeds, test window 2022-2025)", 2)
    rows = build_results_rows()
    if rows:
        add_results_table(doc, rows)
    add_para(
        doc,
        "Reference figures: reports/generated/charts/final_value_comparison.png, "
        "equity_curve_comparison.png and uncertainty_signal.png.",
        italic=True,
    )

    add_heading(doc, "70-ticker test universe — Phase-1 robustness", 2)
    add_para(
        doc,
        "The Phase-1 robustness study runs the same four-agent comparison on a "
        "70-ticker diversified-equity test universe (41 single-name US large-cap "
        "equities spanning technology, payments and financial services, "
        "healthcare, consumer and industrials, plus 29 exchange-traded funds "
        "covering broad-market indices, sector SPDRs, dividend ETFs, thematic "
        "exposures and commodity funds) on the same 2022–2025 test window with "
        "the same metric definitions. The aggregate result is summarised below; "
        "the full per-ticker table is in Appendix B of the main dissertation.",
    )
    cs_table = doc.add_table(rows=1, cols=4)
    cs_table.style = "Light Grid Accent 1"
    cs_hdr = cs_table.rows[0].cells
    for i, h in enumerate(["Strategy", "Mean terminal value", "Mean Sharpe", "Mean Max-DD"]):
        cs_hdr[i].text = h
        for r in cs_hdr[i].paragraphs[0].runs:
            r.bold = True
    for label, final, sharpe, mdd in [
        ("Baseline PPO (no uncertainty)", "$989,430", "−0.23", "0.033"),
        ("Probabilistic PPO (this work)", "$1,998,817", "+0.60", "0.225"),
        ("Manual 5 % trailing stop", "$1,531,163", "+0.36", "0.305"),
        ("Passive buy-and-hold", "$2,099,838", "+0.54", "0.370"),
    ]:
        row = cs_table.add_row().cells
        row[0].text = label
        row[1].text = final
        row[2].text = sharpe
        row[3].text = mdd
        if "Probabilistic" in label:
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.bold = True
    add_para(doc, "Headline findings on the 70-ticker test universe:", bold=True)
    add_bullets(doc, [
        "Drawdown reduced versus passive buy-and-hold on 70 of 70 tickers (100 % of the universe), with an average reduction of 14.5 percentage points (mean drawdown cut by 39 % in relative terms — from 37.0 % to 22.5 %). This is the strongest single number in the dissertation.",
        "Probabilistic agent beat the manually-tuned 5 % trailing stop on 61 of 70 tickers (87 %) in terminal value, and on essentially every ticker in Sharpe ratio — the empirical answer to the previous-meeting question on whether the AI agent beats a manually-tuned stop-loss alternative.",
        "Cost in mean terminal value versus passive buy-and-hold: roughly 5 % give-up in mean upside in exchange for the 39 % reduction in mean drawdown above. This is exactly the trade an institutional drawdown-mandated investor runs every quarter.",
        "Where the agent loses (45 of 70 tickers in terminal value, all winning on drawdown), the losses cluster in two diagnosable regimes: persistent, low-uncertainty bull-market trends in single names (NVDA, AVGO, LLY) where the uncertainty-guard's caution costs the right tail, and very-low-drawdown defensives (JNJ, MCD, SCHD, GLD) where there is essentially nothing for a drawdown overlay to add. Sector-aware uncertainty-quantile calibration is the targeted Phase-2 fix.",
    ])

    add_heading(doc, "How to read these numbers", 2)
    add_bullets(doc, [
        "The headline criterion is the joint of Sharpe ratio and drawdown control, not either half alone. Meeting either half on its own is trivial: an all-cash policy achieves perfect preservation with zero return, and a return-only policy ignores the constraint entirely. On the 70-ticker test universe the probabilistic agent meets the joint constraint: it controls drawdown on 100 % of the universe (mean drawdown 22.5 % vs buy-and-hold's 37.0 %) and earns a higher mean Sharpe than passive buy-and-hold (+0.60 vs +0.54). The baseline PPO meets neither half: it ends roughly where it started with a slightly negative Sharpe.",
        "The manually-tuned trailing-stop comparator (5 % stop with 20/50-day moving-average re-entry) is the directly measured manual alternative. The probabilistic agent beats it on 87 % of the 70-ticker universe in terminal value and on essentially every ticker in Sharpe. This is a directly-measured rather than asserted answer to the previous-meeting question on whether the AI agent beats a manually-tuned stop-loss alternative.",
        "Max drawdown on the baseline looks small only because the baseline barely compounds in the first place. Both terminal preservation and path preservation (1 − MDD) are reported in the dissertation so the reader can apply whichever definition matches their mandate.",
        "These numbers are provisional Phase-1 evidence. The Phase-2 extended grid (10 seeds × 50 000 timesteps × 4 walk-forward folds × 16 bootstrap paths × 70 tickers) on the Colab GPU runtime will tighten the seed-variability bands and provide out-of-time confirmation across the full universe.",
    ])

    add_heading(doc, "Reproducibility", 2)
    p = doc.add_paragraph()
    run = p.add_run(
        "python experiments/run_baseline.py\n"
        "python experiments/run_probabilistic_agent.py\n"
        "python experiments/run_benchmarks.py\n"
        "python experiments/run_rule_baselines.py\n"
        "python reports/generate_dissertation_report.py\n"
        "python reports/build_supervisor_pack.py\n"
        "python reports/plot_dissertation_visuals.py"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    add_para(
        doc,
        "Artifacts land in experiments/results/ and reports/generated/. The full "
        "source is on GitHub at TheFinix13/Dissertation_Sample_Project, with the "
        "walkthrough notebook (Dissertation_Walkthrough.ipynb) as the single entry "
        "point for someone reading the project for the first time.",
    )

    page_break(doc)

    # ----- Future plan -----
    add_heading(doc, "Future plan", 1)
    add_para(
        doc,
        "The phasing below maps onto the form's structure and reflects the revised "
        "direction agreed with the supervisor. Each scheduled task is tied to a "
        "milestone with a target date. Milestones are also tied to objectives O1 to O4.",
    )
    add_plan_table(doc, [
        (
            "May 2026 (remaining)",
            "Reframe the dissertation around drawdown-constrained risk-adjusted return "
            "(Abstract, Chapter 1, interim review). Add a Finance / Risk-Management "
            "Background sub-chapter covering MV, CVaR/ES, drawdown measures and "
            "Sortino, with notation explained. Implement the rule-based stop-loss "
            "baseline as a third comparator. Expand the test universe from a single "
            "index to a 70-ticker diversified-equity universe and re-run the four-"
            "agent comparison at the Phase-1 budget. Add the Section 5.5.1 extended seed-"
            "stability check on a representative eight-ticker sub-universe.",
            "M1: rule-based baseline checked in and reported alongside the AI "
            "agents (mid-May, done). M2: 70-ticker robustness study + extended "
            "seed-stability check on representative sub-universe (end of May, done).",
        ),
        (
            "June 2026 (4 weeks)",
            "Phase-2 extended grid on the full 70-ticker universe at extended "
            "budget (10 seeds × 50 000 timesteps × 4 walk-forward folds × 16 "
            "bootstrap paths) on Colab GPU runtime. Sector-aware uncertainty-"
            "quantile calibration (replace single global threshold with per-"
            "sector or per-regime threshold). Begin Chapter 2 (Background) and "
            "Chapter 3 (Methodology) full drafts.",
            "M3: full 70-ticker × 4-fold × 10-seed × 50k-step extended grid (mid-"
            "June). M4: sector-aware calibration ablation + Chapter 2 and Chapter "
            "3 first drafts (end of June).",
        ),
        (
            "July 2026 (4-6 weeks)",
            "Sensitivity sweep on the uncertainty threshold, minimum scale, and max "
            "trade fraction. Block-bootstrap data augmentation (Politis and Romano, "
            "1994) to expand the effective training set. Locked final results table. "
            "Draft Chapter 5 (Results) and Chapter 1 (Introduction).",
            "M5: sensitivity and bootstrap results locked (mid-July). M6: Chapters 1, "
            "2, 3 and 5 first drafts (end of July).",
        ),
        (
            "August 2026 (4 weeks)",
            "Start the paper-trading shadow run via Alpaca early in the month and let "
            "it accumulate two weeks of out-of-sample profit-and-loss. Write Chapter "
            "6 (Discussion) and Chapter 7 (Conclusion). Polish figures, integrate "
            "supervisor feedback, finalise the dissertation. Code changes from this "
            "point are bug-fix only.",
            "M7: paper-trading shadow run started (early August). M8: full draft to "
            "supervisor (mid-August). M9: paper-trading PnL added to results chapter "
            "(third week of August). M10: submission-ready version (end of August).",
        ),
        (
            "September 2026",
            "Submit by 1 September 2026. Viva preparation: slide deck (no more than "
            "twelve slides, no more than twenty minutes per the project handbook), "
            "demo of the reproducible pipeline, pre-emptive question and answer "
            "rehearsal using reports/templates/viva_qa_notes.md.",
            "M11: viva-ready presentation and demo by viva date.",
        ),
    ])

    add_heading(doc, "Risks and mitigations", 2)
    add_bullets(doc, [
        "Compute time. Current runs are CPU-friendly (10 000 PPO timesteps, three seeds). The multi-ticker, walk-forward, ablation and ten-seed grid is larger but still CPU-tractable; runs will be batched overnight and partial-grid results accepted for any interim deliverable. Where the grid is too heavy for the local machine the corresponding runner can be lifted onto a Google Colab GPU runtime without code changes.",
        "Data-API drift. yfinance occasionally changes its column shape. The _close_1d helper used by every runner already normalises this, and the protocol pins explicit dates so a re-pull stays comparable.",
        "Result fragility. The Phase-1 numbers may move under the multi-ticker, walk-forward and ablation work. To guard against over-claiming, results will be reported as median and inter-quartile range across at least ten seeds and across tickers, evaluated on multiple sliding test windows (walk-forward) rather than a single one, and any case where the probabilistic variant fails to beat the rule-based stop-loss comparator or buy-and-hold will be called out explicitly.",
        "Paper-trading dependency. The Alpaca shadow run depends on a working brokerage account and stable market hours. If the API is unavailable for any portion of August, the dissertation will report whatever live profit-and-loss was accumulated up to the cutoff, with the gap explicitly stated.",
    ])

    page_break(doc)

    # ----- Extenuating circumstances -----
    add_heading(doc, "Extenuating circumstances", 1)
    add_para(
        doc,
        "[Student to fill in. Either record \"None to declare\" or describe and "
        "indicate that the personal tutor and student-support services have been "
        "informed. Do not include medical detail in this document.]",
        italic=True,
    )

    # ----- Self-tick -----
    add_heading(doc, "Indicative project hours and progress", 1)
    add_para(
        doc,
        "Self-assessment of the first one hundred hours allocated to the project. "
        "The student should tick exactly one of the following:",
    )
    add_bullets(doc, [
        "[ ]  The work has exceeded the first 100 hours of time allocated.",
        "[X] The work has sufficiently met the first 100 hours.",
        "[ ]  The majority of the first 100 hours have been completed but some time has been lost and will be made up.",
        "[ ]  Engagement in the project has been insufficient and progress is of concern.",
    ])
    add_para(
        doc,
        "Justification: the reproducible Phase-0 and Phase-1 pipeline, the protocol "
        "document, the baseline and probabilistic agents, the benchmarks, the "
        "rule-based stop-loss comparator and the generated supervisor pack together "
        "support the second tick above.",
        italic=True,
    )

    page_break(doc)

    # ----- Supervisor section -----
    add_heading(doc, "Supervisor's assessment", 1)
    add_para(
        doc,
        "The boxes below are reserved for the supervisor's written feedback. The "
        "student-completed sections of this document are intended to give the "
        "supervisor enough material to assess each item.",
        italic=True,
    )
    add_supervisor_box(doc, "Comments on the project plan and on the student's progress to date:")
    add_supervisor_box(doc, "Comments on the literature review and the framing of the problem:")
    add_supervisor_box(doc, "Comments on the technical progress and the experimental protocol:")
    add_supervisor_box(doc, "Recommendations for the remainder of the project:")
    add_supervisor_box(doc, "Other comments (optional):")

    # ----- Save -----
    out = EXPORTS / "InterimReview.docx"
    doc.save(out)
    print(f"Wrote: {out}")
    return out


if __name__ == "__main__":
    build()
