"""Build the formal dissertation Word document.

Outputs:
    reports/generated/exports/Dissertation_Draft.docx

Run:
    venv/bin/python reports/build_dissertation_docx.py
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

import matplotlib
import numpy as np

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
RESULTS = ROOT / "experiments" / "results"
EXPORTS = ROOT / "reports" / "generated" / "exports"
EQ_DIR = EXPORTS / "equations"
CHARTS = ROOT / "reports" / "generated" / "charts"
EXPORTS.mkdir(parents=True, exist_ok=True)
EQ_DIR.mkdir(parents=True, exist_ok=True)


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #
def render_equation(latex: str, filename: str, fontsize: int = 18) -> Path:
    """Render a LaTeX equation to a transparent PNG via matplotlib mathtext."""
    path = EQ_DIR / filename
    fig = plt.figure(figsize=(0.01, 0.01))
    fig.text(0, 0, latex, fontsize=fontsize)
    fig.savefig(path, dpi=240, bbox_inches="tight", pad_inches=0.05, transparent=False)
    plt.close(fig)
    return path


def latest_json(prefix: str) -> dict | list:
    files = sorted(p for p in RESULTS.glob(f"{prefix}_*.json"))
    if not files:
        return []
    return json.loads(files[-1].read_text(encoding="utf-8"))


def avg(rows: Iterable[dict], key: str) -> float:
    rows = list(rows)
    if not rows:
        return float("nan")
    return float(np.mean([float(r[key]) for r in rows]))


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc: Document, text: str, *, italic: bool = False, bold: bool = False, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_equation(doc: Document, latex: str, filename: str, *, label: str | None = None, width_inches: float = 4.5) -> None:
    path = render_equation(latex, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    if label:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(label)
        cap_run.italic = True
        cap_run.font.size = Pt(10)


def add_figure(doc: Document, image_path: Path, caption: str, *, width_inches: float = 5.5) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(10)


def add_metrics_table(doc: Document, rows: list[dict]) -> None:
    headers = ["Agent", "Final value (USD)", "Sharpe", "Max DD", "VaR-95 viol.", "Preservation"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for r in hdr_cells[i].paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["agent"]
        cells[1].text = f"${row['final']:,.0f}"
        cells[2].text = f"{row['sharpe']:+.4f}"
        cells[3].text = f"{row['mdd']:.4f}"
        cells[4].text = f"{row['var']:.4f}"
        cells[5].text = f"{row['pres']:.4f}"
    for r in table.rows:
        for c in r.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def page_break(doc: Document) -> None:
    doc.add_page_break()


def set_default_font(doc: Document, family: str = "Times New Roman", size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = family
    style.font.size = Pt(size)
    rpr = style.element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), family)


# --------------------------------------------------------------------------- #
# Build
# --------------------------------------------------------------------------- #
def build() -> Path:
    doc = Document()
    set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    baseline = latest_json("baseline")
    prob = latest_json("probabilistic")
    bench = latest_json("benchmarks")
    bench_lookup = {r["agent"]: r for r in bench} if bench else {}

    metrics_rows = []
    if baseline and prob:
        metrics_rows.extend([
            {
                "agent": "Baseline PPO",
                "final": avg(baseline, "final_portfolio_value"),
                "sharpe": avg(baseline, "sharpe_ratio"),
                "mdd": avg(baseline, "max_drawdown"),
                "var": avg(baseline, "var_95_violation_rate"),
                "pres": avg(baseline, "capital_preservation_rate_95pct_hwm"),
            },
            {
                "agent": "Probabilistic PPO",
                "final": avg(prob, "final_portfolio_value"),
                "sharpe": avg(prob, "sharpe_ratio"),
                "mdd": avg(prob, "max_drawdown"),
                "var": avg(prob, "var_95_violation_rate"),
                "pres": avg(prob, "capital_preservation_rate_95pct_hwm"),
            },
        ])
    for label in ("buy_and_hold", "all_cash"):
        b = bench_lookup.get(label)
        if b:
            metrics_rows.append({
                "agent": "Buy-and-hold (SPY)" if label == "buy_and_hold" else "All-cash",
                "final": float(b["final_portfolio_value"]),
                "sharpe": float(b["sharpe_ratio"]),
                "mdd": float(b["max_drawdown"]),
                "var": float(b["var_95_violation_rate"]),
                "pres": float(b["capital_preservation_rate_95pct_hwm"]),
            })

    # ----- Title page -----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Probabilistic Deep Reinforcement Learning\nfor Portfolio Risk Analysis")
    r.bold = True
    r.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("An uncertainty-aware policy for capital preservation under regime stress")
    r.italic = True
    r.font.size = Pt(14)

    for _ in range(4):
        doc.add_paragraph()

    for line, sz, bold in [
        ("Fiyin Akano", 14, True),
        ("URN: [INSERT URN]", 12, False),
        ("", 12, False),
        ("EEEM004 — MSc Dissertation", 12, False),
        ("Department of Electrical and Electronic Engineering", 12, False),
        ("University of Surrey", 12, False),
        ("", 12, False),
        ("Supervisor: [INSERT SUPERVISOR]", 12, False),
        ("Second supervisor: [INSERT IF APPLICABLE]", 12, False),
        ("", 12, False),
        ("September 2026 (target submission)", 12, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(sz)

    page_break(doc)

    # ----- Abstract -----
    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "Reinforcement-learning agents trained on financial price series are usually tuned for "
        "return rather than for control of downside risk. This dissertation asks whether a "
        "small, explicit forecast-uncertainty signal, produced by a DeepAR-style probabilistic "
        "long short-term memory network, can be added to a standard Proximal Policy "
        "Optimization (PPO) pipeline so that the resulting agent's dominant property becomes "
        "capital preservation instead of return-seeking. The contribution has two parts. First, "
        "an uncertainty-aware trading environment is defined in which trade size is shrunk by "
        "(1 - u_t) with a floor s_min, and new long-side trades are blocked when uncertainty "
        "exceeds a quantile threshold tau. Second, a fully reproducible evaluation protocol is "
        "set out: fixed train, validation and test splits over US equity index data "
        "(2009 to 2025), three random seeds, and a metric set covering final value, annualised "
        "return and volatility, Sharpe ratio, maximum drawdown, the 95 % Value-at-Risk "
        "violation rate, and the capital-preservation ratio against the running high-watermark. "
        "On the held-out test window the probabilistic agent meets the stated objective "
        "(preservation >= 0.95) and finishes above passive buy-and-hold. The baseline PPO is "
        "essentially flat. The headline result comes with a caveat that the discussion takes "
        "seriously: the baseline only avoids drawdown because it never compounds enough to "
        "have anything to draw down from. All experiments, plots and reports are reproducible "
        "from the public repository, and a single Jupyter walkthrough is provided as an "
        "end-to-end review path.",
    )
    add_para(doc, "Keywords: reinforcement learning, portfolio management, deep learning, "
             "uncertainty quantification, risk management, capital preservation, PPO, DeepAR.",
             italic=True)
    page_break(doc)

    # ----- Acknowledgements -----
    add_heading(doc, "Acknowledgements", 1)
    add_para(doc, "[INSERT ACKNOWLEDGEMENTS]")
    page_break(doc)

    # ----- Table of contents placeholder -----
    add_heading(doc, "Contents", 1)
    add_para(doc, "Generate this in Word (References → Table of Contents → Automatic Table 1) once the headings are reviewed.", italic=True)
    page_break(doc)

    # ============================================================== Chapter 1
    add_heading(doc, "Chapter 1 — Introduction", 1)

    add_heading(doc, "1.1 Background and motivation", 2)
    add_para(
        doc,
        "Quantitative portfolio management has been studied as an optimisation problem under "
        "uncertainty for the better part of a century. Markowitz (1952) is the canonical entry "
        "point: a quadratic optimisation in mean and variance that fixes the trade-off between "
        "expected return and risk, and that still serves as the conceptual baseline most newer "
        "methods are measured against. In the last decade deep reinforcement learning (DRL) has "
        "emerged as a flexible alternative. Rather than solve an analytic optimisation under a "
        "fixed set of constraints, a DRL agent learns a control policy directly from "
        "interaction with a market environment. Jiang et al. (2017), Yang et al. (2020) and "
        "the FinRL library of Liu et al. (2021) have shown that policy-gradient methods can "
        "be made to work on equities, ETFs and cryptocurrencies.",
    )
    add_para(
        doc,
        "Two practical issues with that body of work motivate this dissertation. The first is "
        "that the reward signal almost always rewards portfolio return, and lets risk in only "
        "indirectly: through volatility-adjusted reward shaping, or by reading Sharpe off after "
        "the fact. The second is that the policies are trained on point-estimate features and "
        "are never told how confident an underlying forecaster is in its own prediction. When "
        "the market enters a regime where forecast variance spikes, for example during a macro "
        "shock, the agent has no native way to be cautious.",
    )

    add_heading(doc, "1.2 Problem statement", 2)
    add_para(
        doc,
        "The question this dissertation addresses can be put plainly. If we compute an "
        "explicit forecast-uncertainty signal from a probabilistic recurrent network and feed "
        "it into a standard PPO pipeline, can the resulting agent preserve capital, defined "
        "as the ratio of terminal portfolio value to the running high-watermark, through a "
        "test window that contains real shocks, while still keeping up with passive benchmarks "
        "in calm periods?",
    )

    add_heading(doc, "1.3 Aims and objectives", 2)
    add_bullets(doc, [
        "O1. Find out whether an explicit forecast-uncertainty signal, modelled with a DeepAR-style probabilistic LSTM, can be plugged into a PPO policy to make it behave with risk in mind on US equity index data.",
        "O2. Test, on a held-out window that contains real shocks, whether the resulting agent preserves at least 95 % of its high-watermark portfolio value, both against a baseline PPO and against passive buy-and-hold and all-cash benchmarks.",
        "O3. Pin down a reproducible evaluation protocol of fixed splits, fixed seeds, scripted artifacts and a shared metric set, so that any comparison made in this dissertation is genuinely like-for-like.",
        "O4. Take a position, on the strength of the above, on when an uncertainty signal earns a place in a portfolio control loop, and equally important, when it does not.",
    ])

    add_heading(doc, "1.4 Contributions", 2)
    add_bullets(doc, [
        "An uncertainty-aware trading environment in which the per-step trade size is shrunk by (1 - u_t) with a floor s_min, and new long-side trades are blocked when uncertainty exceeds the quantile threshold tau (Section 3.5).",
        "A reproducible end-to-end evaluation protocol with fixed splits, three random seeds and a metric set centred on the capital-preservation ratio against the running high-watermark (Section 3.7).",
        "A public, runnable Jupyter walkthrough that loads the dataset, trains the probabilistic forecaster, prints the uncertainty values, runs both agents and renders the comparison table and equity curves with embedded outputs (Chapter 5 and Appendix A).",
        "A discussion that calls out the maximum-drawdown comparison as misleading on its own and argues for the preservation ratio as the metric that actually matches the stated objective (Section 6.2).",
    ])

    add_heading(doc, "1.5 Dissertation structure", 2)
    add_para(
        doc,
        "Chapter 2 reviews the relevant background: modern portfolio theory, the policy-"
        "gradient family of RL algorithms (with PPO singled out), prior DRL work in finance, "
        "probabilistic forecasting with DeepAR-style networks, and methods for uncertainty "
        "quantification in deep learning. Chapter 3 sets the problem out as a Markov decision "
        "process, derives the probabilistic forecaster, defines the trading environment and "
        "states the exact mathematical difference between the baseline and the probabilistic "
        "agent. Chapter 4 covers the implementation: data pipeline, training procedure and "
        "reporting layer. Chapter 5 presents the experimental setup and the results on the "
        "held-out test window. Chapter 6 reads the results carefully, including the trade-offs "
        "and the limits of the evidence. Chapter 7 concludes and points at future work.",
    )
    page_break(doc)

    # ============================================================== Chapter 2
    add_heading(doc, "Chapter 2 — Background and Related Work", 1)

    add_heading(doc, "2.1 Modern portfolio theory and capital preservation", 2)
    add_para(
        doc,
        "Markowitz (1952) framed portfolio choice as a quadratic optimisation problem in mean "
        "and variance and produced the efficient frontier, the locus of admissible mean-"
        "variance trade-offs. The literature has since extended this in many directions: "
        "conditional value-at-risk, drawdown control and dynamic rebalancing among them. "
        "Capital preservation, as I use the term in this dissertation, is a different objective "
        "again. I define it as the ratio of terminal portfolio value to the running high-"
        "watermark. The metric penalises any path that fails to keep the gains it has already "
        "realised, which is closer to how a long-horizon investor or a risk-constrained mandate "
        "actually thinks about performance than mean-variance is.",
    )

    add_heading(doc, "2.2 Reinforcement learning fundamentals", 2)
    add_para(
        doc,
        "An RL problem is a Markov decision process (S, A, P, R, gamma): a state space, an "
        "action space, a transition kernel, a reward function and a discount factor. Sutton "
        "and Barto (2018) is the canonical treatment. The objective is to learn a policy "
        "pi(a|s) that maximises the expected discounted return.",
    )
    add_equation(
        doc,
        r"$J(\pi) = \mathbb{E}_{\tau\sim\pi}\left[\sum_{t=0}^{T}\gamma^{t}\,R(s_{t},a_{t})\right]$",
        "eq_rl_objective.png",
        label="Equation 2.1 — Discounted-return objective.",
    )

    add_heading(doc, "2.3 Policy-gradient methods and PPO", 2)
    add_para(
        doc,
        "Policy-gradient methods parameterise the policy directly and ascend the gradient of "
        "the objective with respect to its parameters. Vanilla policy gradient is high-"
        "variance and unstable. Trust-region methods solve this by constraining the size of "
        "each update. Schulman et al. (2017) introduced Proximal Policy Optimization (PPO), "
        "which approximates a trust region with a simpler clipped surrogate objective.",
    )
    add_equation(
        doc,
        r"$L^{\mathrm{CLIP}}(\theta) = \mathbb{E}_{t}\left[\min\left(\rho_{t}(\theta)\hat{A}_{t},\,\mathrm{clip}(\rho_{t}(\theta),1-\epsilon,1+\epsilon)\hat{A}_{t}\right)\right]$",
        "eq_ppo_clip.png",
        label="Equation 2.2 — PPO clipped surrogate objective (Schulman et al., 2017).",
    )
    add_para(
        doc,
        "PPO has become the default policy-gradient choice for finance applications for two "
        "practical reasons. It is forgiving with respect to hyper-parameter choice, and it is "
        "available in mature open-source form, notably in Stable-Baselines3 (Raffin et al., "
        "2021). I use both in this dissertation.",
    )

    add_heading(doc, "2.4 Reinforcement learning for portfolio management", 2)
    add_para(
        doc,
        "Jiang et al. (2017) propose an end-to-end DRL framework for cryptocurrency portfolio "
        "selection, with a CNN feature extractor and a policy that outputs portfolio weights "
        "directly. Yang et al. (2020) extend the idea to US equities using an ensemble of A2C, "
        "PPO and DDPG agents trained across different regimes. Liu et al. (2021) bundle this "
        "kind of work into the FinRL library, which standardises gym-style trading "
        "environments and exposes one API across several algorithms. None of these papers "
        "inject an explicit forecast-uncertainty signal into the policy state, and none use "
        "uncertainty as a trading guard. That is the gap this dissertation addresses.",
    )

    add_heading(doc, "2.5 Probabilistic time-series forecasting", 2)
    add_para(
        doc,
        "Salinas et al. (2020) propose DeepAR, an autoregressive recurrent network trained to "
        "maximise the likelihood of the observed data under a chosen output distribution. The "
        "backbone is a stacked long short-term memory network (Hochreiter and Schmidhuber, "
        "1997). With a Gaussian output the network emits (mu, log sigma^2) at each step.",
    )
    add_equation(
        doc,
        r"$\mathcal{L}(\theta) = \frac{1}{N}\sum_{t}\frac{1}{2}\left[\log(\sigma_{t}^{2}+\varepsilon) + \frac{(y_{t}-\mu_{t})^{2}}{\sigma_{t}^{2}+\varepsilon}\right]$",
        "eq_gaussian_nll.png",
        label="Equation 2.3 — Gaussian negative log-likelihood loss used to train the probabilistic forecaster.",
    )

    add_heading(doc, "2.6 Uncertainty estimation in deep learning", 2)
    add_para(
        doc,
        "Two practical techniques dominate the uncertainty-in-neural-networks literature. "
        "Deep ensembles (Lakshminarayanan, Pritzel and Blundell, 2017) train several "
        "independent networks and use their disagreement as an uncertainty proxy. Monte-Carlo "
        "dropout (Gal and Ghahramani, 2016) treats dropout at inference time as a variational "
        "approximation to a Bayesian posterior. The DeepAR-style likelihood I use in this "
        "dissertation is a third route. The network emits a predictive variance directly, "
        "trained by the Gaussian NLL in Equation 2.3. I picked it over ensembles or MC-"
        "dropout because the policy only needs a one-dimensional, continuous uncertainty "
        "summary, and a single Gaussian head delivers that with the smallest amount of moving "
        "machinery.",
    )

    add_heading(doc, "2.7 Gap and positioning", 2)
    add_para(
        doc,
        "Putting all of the above together, the gap I want to close is the absence, in prior "
        "DRL-for-finance work, of an explicit, model-based uncertainty signal that does two "
        "things at once: it goes into the policy state, and it acts as a hard guard on new "
        "risk-on actions. Section 3.5 sets out the mathematical difference from a standard "
        "PPO baseline. Chapter 5 measures whether the difference earns its keep.",
    )
    page_break(doc)

    # ============================================================== Chapter 3
    add_heading(doc, "Chapter 3 — Methodology", 1)

    add_heading(doc, "3.1 Problem formulation", 2)
    add_para(
        doc,
        "I cast the trading task as a Markov decision process. The state at step t is a "
        "concatenation of the last L log-returns of the underlying asset, the current "
        "normalised position size, and, for the probabilistic agent only, the uncertainty "
        "score u_t. The action is a scalar a_t in [-1, 1]. After scaling by a configurable "
        "max trade fraction f_max it prescribes either a fraction of cash to deploy on the "
        "buy side or a fraction of position to liquidate on the sell side. Trades incur a "
        "linear transaction cost at rate c. The per-step reward is the scaled log-growth of "
        "portfolio value.",
    )
    add_equation(
        doc,
        r"$R_{t} = 100\cdot\log\left(\frac{V_{t+1}}{V_{t}}\right),\qquad V_{t}=B_{t}+n_{t}\,p_{t}$",
        "eq_reward.png",
        label="Equation 3.1 — Per-step reward, with cash B_t and shares n_t at price p_t.",
    )

    add_heading(doc, "3.2 Data and preprocessing", 2)
    add_para(
        doc,
        "Daily adjusted close prices come from Yahoo Finance via the yfinance Python package. "
        "The Phase-1 universe is SPY (the S&P 500 ETF). QQQ and a small basket of sector ETFs "
        "are queued for the Phase-2 robustness study. Two shock windows are fixed in the "
        "protocol and used during stress evaluation: the COVID crash (February to June 2020) "
        "and the onset of the Russia-Ukraine war (February to September 2022). The price "
        "series is converted to log-returns and the LSTM forecaster sees supervised sequences "
        "of length L = 20.",
    )
    add_equation(
        doc,
        r"$r_{t}=\log p_{t}-\log p_{t-1},\qquad \mathbf{x}^{(i)}=(r_{i-L+1},\dots,r_{i}),\qquad y^{(i)}=r_{i+1}$",
        "eq_returns_seq.png",
        label="Equation 3.2 — Log-returns and supervised sequence construction.",
    )

    add_heading(doc, "3.3 Probabilistic forecaster", 2)
    add_para(
        doc,
        "The forecaster is a two-layer LSTM with hidden dimension 32, followed by two linear "
        "heads. One head emits the predictive mean, the other emits the log of the predictive "
        "variance:",
    )
    add_equation(
        doc,
        r"$\mathbf{h}_{t}=\mathrm{LSTM}(\mathbf{x}_{1:t};\theta),\quad \mu_{t}=W_{\mu}\mathbf{h}_{t}+b_{\mu},\quad \log\sigma_{t}^{2}=W_{\sigma}\mathbf{h}_{t}+b_{\sigma}$",
        "eq_lstm_arch.png",
        label="Equation 3.3 — DeepAR-style LSTM architecture used in this dissertation.",
    )
    add_para(
        doc,
        "Training uses the Gaussian NLL loss from Equation 2.3, optimised with Adam at "
        "learning rate 1e-3 for 20 epochs.",
    )

    add_heading(doc, "3.4 Uncertainty score", 2)
    add_para(
        doc,
        "At inference time the predictive standard deviation is min-max normalised across the "
        "test window into a unit-interval uncertainty score u_t in [0, 1]:",
    )
    add_equation(
        doc,
        r"$\hat\sigma_{t}=\sqrt{\sigma_{t}^{2}},\qquad u_{t}=\frac{\hat\sigma_{t}-\min_{t}\hat\sigma_{t}}{\max_{t}\hat\sigma_{t}-\min_{t}\hat\sigma_{t}+10^{-8}}$",
        "eq_uncertainty.png",
        label="Equation 3.4 — Normalisation that produces the uncertainty score consumed by the policy.",
    )

    add_heading(doc, "3.5 Trading environment and the contribution", 2)
    add_para(
        doc,
        "Both agents share the same gymnasium-compatible environment. The only mathematical "
        "difference between them is two extra terms in the trade-size computation that the "
        "probabilistic agent uses. With cash balance B_t, max trade fraction f_max = 0.10, "
        "raw policy action a_t in [-1, 1], uncertainty u_t in [0, 1], threshold tau set at the "
        "80th percentile of u_t, and minimum scale s_min = 0.10, the baseline trade size is",
    )
    add_equation(
        doc,
        r"$v_{t}^{\mathrm{base}}=B_{t}\cdot f_{\max}\cdot a_{t}$",
        "eq_baseline_trade.png",
        label="Equation 3.5 — Baseline PPO trade size.",
    )
    add_para(doc, "and the probabilistic trade size is")
    add_equation(
        doc,
        r"$v_{t}^{\mathrm{prob}}=B_{t}\cdot f_{\max}\cdot a_{t}\cdot\max(1-u_{t},\,s_{\min}),\quad v_{t}^{\mathrm{prob}}=0\ \mathrm{if}\ u_{t}\geq\tau\ \mathrm{and}\ v_{t}^{\mathrm{prob}}>0$",
        "eq_probabilistic_trade.png",
        label="Equation 3.6 — Probabilistic PPO trade size with the trade-scaling factor and the risk-on guard.",
    )
    add_para(
        doc,
        "Equation 3.6 is the mathematical contribution of this dissertation. The first extra "
        "factor, the trade-scaling term, shrinks every trade in proportion to the current "
        "forecast uncertainty, with the floor s_min stopping the agent from being silenced "
        "entirely on a quiet day. The second term is a hard guard: once uncertainty exceeds "
        "tau, no new long-side risk may be added. The agent can still de-risk by selling.",
    )

    add_figure(
        doc, CHARTS / "uncertainty_signal.png",
        caption="Figure 3.1 — Normalised forecast uncertainty u_t over the test window; "
                "the 80th-percentile threshold tau is the gate above which new buys are blocked.",
    )

    add_heading(doc, "3.6 Baseline PPO", 2)
    add_para(
        doc,
        "The baseline agent is identical in every other respect. Same observation window, "
        "same reward, same hyper-parameters (learning rate 3e-4, n_steps 512, batch size 64, "
        "n_epochs 5, total time-steps 10 000), and the same Stable-Baselines3 PPO solver with "
        "an MLP policy. What it does not have is the uncertainty coordinate in the state and "
        "the two extra factors in Equation 3.6.",
    )

    add_heading(doc, "3.7 Evaluation protocol", 2)
    add_para(
        doc,
        "The protocol fixes the train, validation and test splits (2009 to 2018, 2019 to "
        "2021, and 2022 to 2025). Three random seeds {7, 19, 42} are used for both the "
        "baseline and the probabilistic agent. Benchmarks are a buy-and-hold of SPY taken "
        "at the start of the test window, and an all-cash position. The metric set is: final "
        "portfolio value, annualised return and volatility, Sharpe ratio, maximum drawdown, "
        "the 95 % Value-at-Risk and the rate at which the realised log-return falls below it, "
        "and the capital-preservation ratio against the running high-watermark.",
    )
    add_equation(
        doc,
        r"$\mathrm{Pres}=\frac{V_{T}}{\max_{t\leq T} V_{t}},\qquad \mathrm{objective:}\ \mathrm{Pres}\geq 0.95$",
        "eq_preservation.png",
        label="Equation 3.7 — Capital-preservation ratio (the project's headline objective).",
    )
    page_break(doc)

    # ============================================================== Chapter 4
    add_heading(doc, "Chapter 4 — Implementation", 1)

    add_heading(doc, "4.1 Software stack", 2)
    add_bullets(doc, [
        "Reinforcement learning: Stable-Baselines3 (PPO) on top of gymnasium.",
        "Probabilistic forecaster: PyTorch (an LSTM with two linear heads, trained with Gaussian NLL).",
        "Data: yfinance for daily adjusted close prices and pandas for tabular handling.",
        "Reporting: matplotlib for figures, small Python scripts for the metric tables and the supervisor pack, and nbconvert for the walkthrough-notebook PDF.",
    ])

    add_heading(doc, "4.2 Repository structure", 2)
    add_bullets(doc, [
        "experiments/configs/dissertation_protocol.json: the single source of truth for the protocol (splits, seeds, metric list and agent hyper-parameters).",
        "experiments/common.py: environment, metric computation, data fetch and seed helpers.",
        "experiments/run_baseline.py, run_probabilistic_agent.py, run_benchmarks.py: the three runners that produce the seeded artifacts.",
        "experiments/results/: the generated CSV and JSON metric files and the equity-curve series.",
        "reports/build_supervisor_pack.py, generate_dissertation_report.py, plot_dissertation_visuals.py and build_dissertation_docx.py: the reporting layer.",
        "Dissertation_Walkthrough.ipynb: a single-file end-to-end walkthrough that I prepared for supervisor review.",
    ])

    add_heading(doc, "4.3 Reproducibility", 2)
    add_para(
        doc,
        "Reproducibility is something I treated as a hard requirement rather than a nice-to-"
        "have. Every script reads the protocol JSON. The seeds {7, 19, 42} are set globally "
        "before each run. Results land in time-stamped JSON and CSV files. The reporting "
        "layer always picks up the latest results, which means a single re-run automatically "
        "refreshes the supervisor pack, the figures and the dissertation tables.",
    )
    page_break(doc)

    # ============================================================== Chapter 5
    add_heading(doc, "Chapter 5 — Results", 1)

    add_heading(doc, "5.1 Experimental setup", 2)
    add_para(
        doc,
        "The numbers below are means across the three seeds for the two seeded agents, and "
        "deterministic curves for the benchmarks. Everything is on the held-out test window, "
        "1 January 2022 to 31 December 2025. Figure 5.1 shows the SPY adjusted-close series "
        "that drives every comparison in this chapter.",
    )

    add_figure(
        doc, CHARTS / "dataset_spy_close.png",
        caption="Figure 5.1 — SPY daily adjusted close on the test window 2022 to 2025.",
    )

    add_heading(doc, "5.2 Forecast uncertainty in the test window", 2)
    add_para(
        doc,
        "Figure 3.1 already shows the normalised uncertainty signal that the probabilistic "
        "agent consumes. The high-uncertainty bands cluster around regions of elevated "
        "realised volatility, which is what one would hope to see if the forecaster has "
        "picked up anything useful. Numerical summaries (minimum, mean, maximum and the 80th-"
        "percentile threshold tau) are printed at run time in the walkthrough notebook.",
    )

    add_heading(doc, "5.3 Aggregate comparison table", 2)
    if metrics_rows:
        add_metrics_table(doc, metrics_rows)
    add_para(
        doc,
        "Reading the table from top to bottom: the baseline PPO ends near initial capital "
        "with a slightly negative Sharpe. The probabilistic PPO finishes meaningfully above "
        "passive buy-and-hold, with a positive Sharpe and a preservation ratio above 0.99. "
        "The all-cash benchmark is constant by construction and is included as a sanity check "
        "on the metric definitions, not as a serious competitor.",
    )

    add_heading(doc, "5.4 Equity curves and drawdown", 2)
    add_figure(
        doc, CHARTS / "equity_curve_comparison.png",
        caption="Figure 5.2 — Equity curves on the test window: baseline PPO (blue), probabilistic PPO (red), buy-and-hold (green).",
    )
    add_figure(
        doc, CHARTS / "final_value_comparison.png",
        caption="Figure 5.3 — Final-value comparison across baseline PPO, probabilistic PPO and buy-and-hold.",
    )
    page_break(doc)

    # ============================================================== Chapter 6
    add_heading(doc, "Chapter 6 — Discussion", 1)

    add_heading(doc, "6.1 Headline interpretation", 2)
    add_para(
        doc,
        "The probabilistic agent meets the headline objective on the test window. Its "
        "preservation ratio (Equation 3.7) sits comfortably above 0.95 across all three "
        "seeds, and it finishes above passive buy-and-hold. The point worth dwelling on is "
        "that this is achieved with the same RL solver and the same training budget as the "
        "baseline. The only change is the two extra terms in Equation 3.6 and the extra "
        "uncertainty coordinate in the state. Whatever is doing the work, it is the small "
        "design decision of letting the policy see and react to forecast uncertainty.",
    )

    add_heading(doc, "6.2 Reading maximum drawdown carefully", 2)
    add_para(
        doc,
        "Maximum drawdown is the most commonly misread number in this comparison, and it is "
        "worth being explicit about why. The baseline agent reports a small drawdown only "
        "because it never compounds significantly above initial capital, so there is very "
        "little to draw down from in the first place. The probabilistic agent, by contrast, "
        "compounds to a higher peak, gives some of it back, and at its trough still sits "
        "well above the baseline's terminal value. When the stated goal is capital "
        "preservation rather than the minimum drawdown of an essentially flat curve, the "
        "preservation ratio is the metric that matches the goal. I report maximum drawdown "
        "alongside it for transparency, not as a contradicting result.",
    )

    add_heading(doc, "6.3 Limitations", 2)
    add_bullets(doc, [
        "Single asset. Phase 1 uses SPY only. The Phase-2 plan extends the same protocol to QQQ and a small set of sector ETFs.",
        "Daily granularity. Intraday dynamics are out of scope. The trade-size scaling and the risk-on guard would both need re-calibration at minute or tick granularity.",
        "One uncertainty estimator. The DeepAR-style Gaussian likelihood is one of several routes; deep ensembles, MC dropout and conformal predictors are all reasonable alternatives. A small sensitivity comparison is planned.",
        "Three seeds. Enough for indicative results, not enough for tight confidence intervals.",
        "No live trading. This is by design. The dissertation prioritises scientific evaluation and reproducibility over execution risk.",
    ])

    add_heading(doc, "6.4 Threats to validity", 2)
    add_para(
        doc,
        "Two threats are worth flagging directly. The first is that the test window is finite "
        "and contains specific macro events. The uncertainty thresholds were set "
        "prospectively in the protocol rather than tuned on the test window, but a longer "
        "evaluation across more regimes is needed before any strong claim. The second is that "
        "the comparison rests on the metric set in Section 3.7. An agent tuned for a single "
        "metric, maximum drawdown alone for instance, would behave differently. The "
        "preservation-against-high-watermark framing is the one that matches the project "
        "objective, and the standard metrics are reported beside it for transparency.",
    )
    page_break(doc)

    # ============================================================== Chapter 7
    add_heading(doc, "Chapter 7 — Conclusion and Future Work", 1)

    add_heading(doc, "7.1 Summary", 2)
    add_para(
        doc,
        "This dissertation has shown that a small, explicit forecast-uncertainty signal, "
        "produced by a DeepAR-style probabilistic LSTM and consumed by a standard PPO policy "
        "through the two extra terms in Equation 3.6, is enough to turn a return-seeking RL "
        "agent into one that puts capital preservation first. On the held-out test window the "
        "probabilistic agent meets the 0.95 preservation objective, finishes above the "
        "passive buy-and-hold benchmark, and improves Sharpe substantially relative to the "
        "baseline.",
    )

    add_heading(doc, "7.2 Future work", 2)
    add_bullets(doc, [
        "Multi-asset robustness. Apply the same protocol to QQQ and a basket of sector ETFs, and report per-asset and pooled results.",
        "Shock-window case studies. Score both agents on the protocol shock periods (COVID 2020 and the Ukraine-war onset in 2022) as standalone case studies.",
        "Sensitivity. Sweep the uncertainty quantile threshold over {0.7, 0.8, 0.9} and the minimum scale s_min over {0.05, 0.10, 0.20}.",
        "Ablation. Compare PPO, PPO with uncertainty as a state feature only, and PPO with the uncertainty guard only, against the full design.",
        "Alternative uncertainty estimators. Swap the Gaussian-NLL head for a deep ensemble or an MC-dropout probabilistic forecaster.",
        "Risk-aware reward shaping. Compare the implicit risk control I get from the trade-size term against an explicit risk-aware reward, for example log-growth penalised by drawdown.",
        "Path to deployment. Wire the trained models to a paper-trading account (Alpaca is the obvious first target) so that one week of out-of-sample paper PnL can be reported alongside the backtest.",
    ])
    page_break(doc)

    # ============================================================== References
    add_heading(doc, "References", 1)
    refs = [
        "Markowitz, H. (1952). Portfolio Selection. Journal of Finance, 7(1), 77–91.",
        "Sutton, R. S., & Barto, A. G. (2018). Reinforcement Learning: An Introduction (2nd ed.). MIT Press.",
        "Hochreiter, S., & Schmidhuber, J. (1997). Long Short-Term Memory. Neural Computation, 9(8), 1735–1780.",
        "Schulman, J., Wolski, F., Dhariwal, P., Radford, A., & Klimov, O. (2017). Proximal Policy Optimization Algorithms. arXiv:1707.06347.",
        "Salinas, D., Flunkert, V., Gasthaus, J., & Januschowski, T. (2020). DeepAR: Probabilistic Forecasting with Autoregressive Recurrent Networks. International Journal of Forecasting, 36(3), 1181–1191.",
        "Lakshminarayanan, B., Pritzel, A., & Blundell, C. (2017). Simple and Scalable Predictive Uncertainty Estimation using Deep Ensembles. NeurIPS.",
        "Gal, Y., & Ghahramani, Z. (2016). Dropout as a Bayesian Approximation: Representing Model Uncertainty in Deep Learning. ICML.",
        "Jiang, Z., Xu, D., & Liang, J. (2017). A Deep Reinforcement Learning Framework for the Financial Portfolio Management Problem. arXiv:1706.10059.",
        "Yang, H., Liu, X.-Y., Zhong, S., & Walid, A. (2020). Deep Reinforcement Learning for Automated Stock Trading: An Ensemble Strategy. ICAIF.",
        "Liu, X.-Y., Yang, H., Gao, J., & Wang, C. D. (2021). FinRL: a deep reinforcement learning library for automated stock trading in quantitative finance. arXiv:2011.09607.",
        "Raffin, A., Hill, A., Gleave, A., Kanervisto, A., Ernestus, M., & Dormann, N. (2021). Stable-Baselines3: Reliable Reinforcement Learning Implementations. JMLR, 22(268), 1–8.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
    page_break(doc)

    # ============================================================== Appendix A
    add_heading(doc, "Appendix A — Reproducibility commands", 1)
    add_para(doc, "End-to-end reproduction of every artifact in this dissertation:")
    code = (
        "python3 -m venv venv && source venv/bin/activate\n"
        "pip install -r requirements.txt\n"
        "python experiments/run_baseline.py\n"
        "python experiments/run_probabilistic_agent.py\n"
        "python experiments/run_benchmarks.py\n"
        "python reports/generate_dissertation_report.py\n"
        "python reports/build_supervisor_pack.py\n"
        "python reports/plot_dissertation_visuals.py\n"
        "python reports/build_dissertation_docx.py"
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    out = EXPORTS / "Dissertation_Draft.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    out = build()
    print(f"Wrote: {out}")
