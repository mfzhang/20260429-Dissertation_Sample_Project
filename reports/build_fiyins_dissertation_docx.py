"""Build the standalone personal-portfolio Word document.

Outputs:
    reports/generated/exports/Fiyins_Dissertation.docx

This is the canonical home for everything to do with Fiyin's actual brokerage
portfolio: the headline numbers, all 70 per-ticker rows, the visualisations,
the wins-and-losses commentary, the practical implications and the caveats.

It is intentionally kept separate from ``Main_Dissertation_Draft.docx`` so that:

* the academic dissertation can be evaluated on its formal eight-ticker basket
  and four-ticker walk-forward grid alone, without any chosen-book contamination;
* the personal portfolio results are still written up in the same level of
  detail, but in a register a non-quantitative reader (or a finance supervisor
  who wants the practical "so what?") can actually read in one sitting.

Run order (the case-study tables / charts must already exist on disk):

    venv/bin/python reports/build_fiyins_case_study.py            # tables + PNGs
    venv/bin/python reports/build_fiyins_dissertation_docx.py     # this script
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from build_fiyins_case_study import _load_protocol_counts, build_tables  # type: ignore

ROOT = Path(__file__).resolve().parent.parent
EXPORTS = ROOT / "reports" / "generated" / "exports"
CHARTS = ROOT / "reports" / "generated" / "charts"
EXPORTS.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------


def set_default_font(doc: Document, family: str = "Calibri", size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = family
    style.font.size = Pt(size)
    rpr = style.element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), family)


def set_margins(doc: Document, cm: float = 2.2) -> None:
    for section in doc.sections:
        section.top_margin = Cm(cm)
        section.bottom_margin = Cm(cm)
        section.left_margin = Cm(cm)
        section.right_margin = Cm(cm)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold


def add_callout(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(label + " ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    body_run = p.add_run(body)
    body_run.italic = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_figure(doc: Document, image_path: Path, caption: str, *, width_inches: float = 6.5) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(10)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def fmt_money(x):
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "—"
    return f"${x:,.0f}"


def fmt_money_m(x):
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "—"
    return f"${x / 1_000_000:.2f}M"


def fmt_pct(x):
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "—"
    return f"{x * 100:.1f}%"


def fmt_sharpe(x):
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "—"
    return f"{x:+.2f}"


# ---------------------------------------------------------------------------
# Block builders
# ---------------------------------------------------------------------------


def add_headline_table(doc: Document, agg: dict) -> None:
    headers = ["Strategy", "Mean terminal value", "Mean Sharpe", "Mean Max-DD"]
    rows = [
        ("Passive buy-and-hold", agg["bh_final_mean"], agg["bh_sharpe_mean"], agg["bh_mdd_mean"]),
        ("Manual 5% trailing stop-loss", agg["r5_final_mean"], agg["r5_sharpe_mean"], agg["r5_mdd_mean"]),
        ("Baseline PPO (no uncertainty signal)", agg["base_final_mean"], agg["base_sharpe_mean"], None),
        ("Probabilistic PPO (this dissertation)", agg["prob_final_mean"], agg["prob_sharpe_mean"], agg["prob_mdd_mean"]),
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for label, final, sharpe, mdd in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = fmt_money(final)
        cells[2].text = fmt_sharpe(sharpe)
        cells[3].text = fmt_pct(mdd) if mdd is not None else "—"
        if "Probabilistic" in label:
            for c in cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_per_ticker_table(doc: Document, rows: list[dict]) -> None:
    headers = [
        "Ticker", "B&H final", "B&H DD",
        "Prob final", "Prob DD",
        "Stop-5% final",
        "Prob>B&H?", "Prob>Stop?",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = r["ticker"]
        cells[1].text = fmt_money(r["bh_final"])
        cells[2].text = fmt_pct(r["bh_mdd"])
        cells[3].text = fmt_money(r["prob_final"])
        cells[4].text = fmt_pct(r["prob_mdd"])
        cells[5].text = fmt_money(r["r5_final"])
        cells[6].text = (r["prob_vs_bh"] or "").upper()
        cells[7].text = (r["prob_vs_r5"] or "").upper()
        for c in cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------


def build() -> Path:
    data = build_tables()
    rows = data["rows"]
    agg = data["aggregate"]
    counts = _load_protocol_counts()
    n_book = counts["n_portfolio"] or agg["n_tickers"]
    n_stocks = counts["n_stocks"]
    n_etfs = counts["n_etfs"]
    per_ticker_capital = 1_000_000

    # Derived
    dd_wins = sum(
        1 for r in rows
        if r["prob_mdd"] is not None and r["bh_mdd"] is not None and r["prob_mdd"] < r["bh_mdd"]
    )
    dd_total = sum(
        1 for r in rows
        if r["prob_mdd"] is not None and r["bh_mdd"] is not None
    )
    avg_dd_reduction = float(np.mean([
        (r["bh_mdd"] - r["prob_mdd"]) * 100
        for r in rows
        if r["prob_mdd"] is not None and r["bh_mdd"] is not None
    ])) if dd_total else 0.0
    terminal_giveup_pct = (
        (agg["bh_final_mean"] - agg["prob_final_mean"]) / agg["bh_final_mean"] * 100
        if agg["bh_final_mean"] else 0.0
    )
    dd_relative_reduction_pct = (
        (1 - agg["prob_mdd_mean"] / agg["bh_mdd_mean"]) * 100
        if agg["bh_mdd_mean"] else 0.0
    )
    win_rate_r5_pct = (
        (agg["prob_wins_vs_r5"] / agg["n_tickers"]) * 100
        if agg["n_tickers"] else 0.0
    )

    doc = Document()
    set_default_font(doc)
    set_margins(doc, cm=2.2)

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------
    cover_top = doc.add_paragraph()
    cover_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_top_run = cover_top.add_run("Fiyin's Portfolio Dissertation")
    cover_top_run.bold = True
    cover_top_run.font.size = Pt(26)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        "A drawdown-controlled probabilistic-RL policy applied to a 70-ticker\n"
        "real brokerage book — what it did, why it did it, and what it means."
    )
    sub_run.italic = True
    sub_run.font.size = Pt(13)

    for _ in range(2):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        "Companion document to the EEEM004 dissertation\n"
        f"Fiyinfoluwa B. Akano · URN 6962514 · {datetime.now(timezone.utc).strftime('%B %Y')}"
    )
    meta_run.font.size = Pt(11)

    doc.add_paragraph()

    add_callout(
        doc,
        "Plain-English note.",
        (
            "This document is meant to be readable by a finance person who is not also a machine-"
            "learning researcher, and by a machine-learning researcher who is not also a finance person. "
            "I have tried to keep the language simple and to define every finance term and every modelling "
            "term the first time it shows up. Where I had to choose between sounding precise and sounding "
            "human, I chose human. The numbers are the same numbers that are in the academic dissertation; "
            "the prose is just a different register."
        ),
    )

    add_callout(
        doc,
        "Important disclaimer.",
        (
            "This document uses publicly available daily price data (via Yahoo Finance) on the "
            f"{n_book} tickers held in my live brokerage accounts as of April 2026. No personal financial "
            "information from any broker — no trade history, no balances, no API access — is processed. The "
            "numbers below come from a backtest of a research policy, not from a real account. Nothing here "
            "is financial advice. The agent's hypothetical positions are not my actual trades."
        ),
    )

    page_break(doc)

    # ------------------------------------------------------------------
    # Table of contents (manual, not auto-generated)
    # ------------------------------------------------------------------
    add_heading(doc, "Contents", 1)
    toc_items = [
        "Executive summary — what happened, in one paragraph",
        "Chapter 1 — Why this document exists, and who it's for",
        "Chapter 2 — The book under test",
        "Chapter 3 — How the test was run, in plain English",
        "Chapter 4 — Headline results",
        "Chapter 5 — Visual summary",
        "Chapter 6 — Per-ticker breakdown",
        "Chapter 7 — Where the agent shone and where it stumbled",
        "Chapter 8 — What this means for someone who actually owns this book",
        "Chapter 9 — Honest caveats",
        "Chapter 10 — What's still on the to-do list",
        "Appendix A — Reproducibility commands",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.6)
    page_break(doc)

    # ------------------------------------------------------------------
    # Executive summary
    # ------------------------------------------------------------------
    add_heading(doc, "Executive summary — what happened, in one paragraph", 1)
    add_para(doc, (
        f"I took the {n_book} tickers I actually own across two retail brokerage accounts — "
        f"{n_stocks} single-name stocks and {n_etfs} exchange-traded funds — and on each one "
        f"I asked four different strategies to manage ${per_ticker_capital:,.0f} of starting capital "
        f"over the four-year window 2022–2025. Three of the strategies are traditional and require "
        f"no AI: passive buy-and-hold, all-cash, and a manually-tuned 5% trailing stop-loss. The "
        f"fourth is the probabilistic reinforcement-learning policy from the academic dissertation: "
        f"it watches a probabilistic forecast of next-day returns, and when the forecast becomes "
        f"unusually uncertain it reduces position size, with a hard rule that the running portfolio "
        f"value must never fall below 95% of its observed all-time high. Across all "
        f"{n_book} tickers, the AI policy reduced peak-to-trough drawdown versus buy-and-hold on "
        f"{dd_wins} of {dd_total} (the average drop fell from {agg['bh_mdd_mean']*100:.0f}% to "
        f"{agg['prob_mdd_mean']*100:.0f}%, an absolute cut of {avg_dd_reduction:.0f} percentage points); "
        f"it gave up roughly {terminal_giveup_pct:.0f}% of the average final value to do so; and it beat "
        f"the manually-tuned stop-loss in terminal value on "
        f"{agg['prob_wins_vs_r5']} of {agg['n_tickers']} tickers ({win_rate_r5_pct:.0f}%) and on essentially "
        f"every ticker on Sharpe ratio. That is the entire story; the rest of this document is the "
        f"detail."
    ))

    page_break(doc)

    # ------------------------------------------------------------------
    # Chapter 1 — why this document
    # ------------------------------------------------------------------
    add_heading(doc, "Chapter 1 — Why this document exists, and who it's for", 1)
    add_heading(doc, "1.1 Why a separate document at all", 2)
    add_para(doc, (
        "The academic EEEM004 dissertation is evaluated on a deliberately neutral asset universe — eight "
        "broad-market and sector US ETFs (SPY, QQQ, IWM, XLK, XLF, XLE, XLV, XLU) — plus a four-fold walk-"
        "forward grid on four of those tickers. The evaluator never gets to see a ticker that the author "
        "had any personal opinion about. That is on purpose: it keeps the comparison apples-to-apples and "
        "removes any suspicion that the result is an artefact of a chosen book."
    ))
    add_para(doc, (
        "But the question Dr Nguyen asked at the last meeting was a different one. He asked whether the "
        "method is useful in finance. The cleanest possible answer to that question is to take the same "
        "method, apply it without modification to a real, heterogeneous, opinionated portfolio that an "
        "actual person actually owns, and report what happened, ticker by ticker — including the cases "
        "where the method got beaten. That is what this document does. It is kept as its own file so the "
        "academic dissertation can be marked on its formal evidence, and a finance reader (or a future "
        "supervisor, or a viva panel) can pick up this companion document if and only if they want the "
        "real-world stress test."
    ))
    add_heading(doc, "1.2 Who should read which", 2)
    add_bullets(doc, [
        "If you only want to know whether the method works on neutral, academic test assets, read "
        "Main_Dissertation_Draft.docx Chapters 5 and 6 and stop there.",
        "If you only want to know whether the method survives contact with a real, opinionated portfolio, "
        "this document is self-contained — you can read it from the executive summary and never need to "
        "open the academic dissertation.",
        "If you are Dr Nguyen, both documents together answer both questions, and you can choose whether "
        "to fold this case study into the academic dissertation as an appendix or leave it as a stand-"
        "alone — the structure here supports either decision without rewriting.",
    ])

    page_break(doc)

    # ------------------------------------------------------------------
    # Chapter 2 — the book
    # ------------------------------------------------------------------
    add_heading(doc, "Chapter 2 — The book under test", 1)
    add_heading(doc, "2.1 What I actually own", 2)
    add_para(doc, (
        f"The book is {n_book} tickers in total, drawn from two live retail brokerage accounts — Yochaa "
        f"for US single-name equities and Bamboo for ETFs. {n_stocks} of those are single-name stocks; "
        f"{n_etfs} are exchange-traded funds. The single names lean US large-cap technology (Apple, "
        f"Microsoft, NVIDIA, Meta, Alphabet, Amazon, Broadcom, AMD, Oracle, Palantir, TSMC), but the "
        f"book is wider than that: it carries payments and financials (Visa, Mastercard, JP Morgan, "
        f"Goldman Sachs, BlackRock, Schwab, Berkshire Hathaway, Royal Bank of Canada, HSBC), defensive "
        f"large-caps (Johnson & Johnson, McDonald's, Costco, Procter, Caterpillar), healthcare and biotech "
        f"(Eli Lilly, Regeneron, Rythm), an industrials sleeve (Boeing, Caterpillar) and a consumer-tech "
        f"sleeve (Netflix, Spotify, T-Mobile US, Tesla)."
    ))
    add_para(doc, (
        f"The {n_etfs} ETFs are deliberately broad: the four big broad-market funds (SPY, VOO, VTI, DIA), "
        f"large-cap growth and value sleeves (SPYG, SPYV, SCHV, SCHG), small caps (IJR), Schwab's "
        f"international (SCHF, SCHB, SCHK), the dividend ETFs (SCHD, VYM, VYMI), the sector SPDRs I run "
        f"as overlays (XLK technology, XLF financials, XLI industrials, XLU utilities), the thematic "
        f"sleeves (TAN solar, COPX copper miners, IBB biotech), and the precious-metals and commodity "
        f"funds (GLD gold, IAU gold, SLV silver, PPLT platinum). There is also an emerging-markets ESG "
        f"sleeve (EEMX). The full list of normalised tickers is in the "
        f"experiments/configs/dissertation_protocol.json file under the named group fiyins_portfolio."
    ))
    add_heading(doc, "2.2 Why this is a hard test", 2)
    add_para(doc, (
        "Three things make this book a genuinely difficult stress test for any risk-control strategy. "
        "First, the heterogeneity: the same overlay has to manage individual mega-cap stocks (which can "
        "move 10% on an earnings day) and broad ETFs (which barely move 1% on an average day) using the "
        "same machinery. Second, the asymmetry: mid-2022 saw deep drawdowns in growth tech (Meta down "
        "70%, Spotify down 70%, Netflix down 70%, Palantir worse) but barely a scratch on utilities, "
        "consumer staples, dividend ETFs and a few defensives. A risk-control overlay that cannot tell "
        "which side of the book is which will under-perform on both. Third, the persistence of the 2023–"
        "2025 AI-led rally — NVIDIA up roughly 8x, Broadcom up several-fold, Palantir up nearly an order "
        "of magnitude — punishes any caution that a risk-control overlay shows during a strong, "
        "low-volatility uptrend. If the overlay is good, it will not pay too much for that caution; if "
        "it is brittle, it will leave a lot of money on the table."
    ))
    add_heading(doc, "2.3 What I'm asking the agent to do for me", 2)
    add_para(doc, (
        "Plainly: stay in markets when the forecast is confident and the trend is up; trim or step "
        "aside when the forecast becomes uncertain and the running portfolio value is approaching 95% of "
        "its prior all-time high. Never accept a peak-to-trough drawdown deeper than 5% on the running "
        "portfolio if it can be helped, but do not become a permanent cash position either. In the "
        "language of the academic dissertation (Section 3.1.5), this is a constrained optimisation problem: "
        "maximise the expected risk-adjusted return subject to a 95% high-watermark drawdown constraint. "
        "Translated into ordinary English, that's exactly what an institutional risk officer, an "
        "endowment manager, or an attentive private investor would ask their book to do."
    ))

    page_break(doc)

    # ------------------------------------------------------------------
    # Chapter 3 — how the test was run
    # ------------------------------------------------------------------
    add_heading(doc, "Chapter 3 — How the test was run, in plain English", 1)
    add_heading(doc, "3.1 The four strategies on the start line", 2)
    add_bullets(doc, [
        "Passive buy-and-hold. Buy each ticker on day 1 of the test window with $1,000,000 of paper money, "
        "do nothing else for four years, mark to market at the end. This is the benchmark every active "
        "strategy has to justify itself against.",
        "Manual 5% trailing stop-loss. The classic textbook risk-management rule: when the price falls "
        "5% from its highest point since you bought, sell everything. Re-enter when the 20-day moving "
        "average crosses above the 50-day moving average (a common momentum re-entry trigger). This is the "
        "strategy a careful retail investor running spreadsheet alerts would actually use by hand.",
        "Baseline PPO. The same reinforcement-learning algorithm as the dissertation's main agent, but "
        "without any uncertainty signal. Included only as a controlled lower bound — it shows what RL "
        "with the same architecture but the wrong inputs can do, so the eventual probabilistic agent's "
        "improvement cannot be confused with a generic RL artefact.",
        "Probabilistic PPO (the dissertation's contribution). Two neural networks. The first is a "
        "DeepAR-style LSTM that produces both a mean forecast and a variance forecast for tomorrow's "
        "return, trained on each ticker's history. The second is the trading policy, also a neural "
        "network, which reads the current state of the portfolio plus the forecast variance and decides "
        "how much to buy, hold or sell. The forecast variance is used both as an input feature and as "
        "a hard guard: when uncertainty is in the top quartile of its historical distribution, the "
        "policy is barred from opening new long positions.",
    ])
    add_heading(doc, "3.2 The clock and the metric set", 2)
    add_para(doc, (
        "Every cell — every (ticker, strategy) pair — runs over the same 2022-01-01 to 2025-12-31 test "
        "window with the same starting capital ($1,000,000), the same data source (Yahoo Finance daily "
        "adjusted-close), the same transaction cost assumption (10 basis points per trade) and the same "
        "metric set: terminal portfolio value, Sharpe ratio, maximum drawdown (the worst peak-to-trough "
        "fall over the window) and the path-preservation ratio (1 minus max drawdown — i.e. the fraction "
        "of the all-time peak that the worst-case drawdown left intact)."
    ))
    add_heading(doc, "3.3 Why this is a fair fight", 2)
    add_para(doc, (
        "The probabilistic agent is not given any advantage that the manual stop-loss is not. Both can "
        "see daily closing prices only — no intraday, no fundamentals, no news. Both pay the same "
        "transaction cost. Both run on each ticker independently with no portfolio-level co-ordination. "
        "Both face exactly the same window. The only difference is that the agent has been allowed to "
        "look at the pre-2022 history (during training) to learn what 'unusual uncertainty' looks like, "
        "while the stop-loss has no equivalent training step — but the stop-loss thresholds (5% trail, "
        "20-day vs 50-day moving-average re-entry) are exactly the conventional textbook values."
    ))

    page_break(doc)

    # ------------------------------------------------------------------
    # Chapter 4 — headline results
    # ------------------------------------------------------------------
    add_heading(doc, "Chapter 4 — Headline results", 1)
    add_para(doc, (
        f"All numbers below are unweighted means across the {agg['n_tickers']}-ticker book on the "
        f"2022–2025 test window, $1,000,000 of starting capital per ticker, three random seeds for the "
        f"two reinforcement-learning agents (median + IQR seed-stability bands at the extended budget "
        f"are reported separately in Section 10.2)."
    ))
    add_headline_table(doc, agg)
    cap_paragraph = doc.add_paragraph()
    cap_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_paragraph.add_run(
        f"Table 4.1 — Aggregate results across the {agg['n_tickers']}-ticker brokerage book."
    )
    cap_run.italic = True
    cap_run.font.size = Pt(10)

    add_heading(doc, "4.1 The drawdown story", 2)
    add_para(doc, (
        f"This is the headline number of the entire dissertation, and the case study delivers it on a "
        f"real book. Across all {dd_total} tickers, the probabilistic agent reduced maximum drawdown "
        f"versus passive buy-and-hold on {dd_wins} of {dd_total} of them — a "
        f"{(dd_wins/dd_total)*100:.0f}% rate. The average drawdown reduction was {avg_dd_reduction:.1f} "
        f"percentage points: a buy-and-hold portfolio that drew down {agg['bh_mdd_mean']*100:.0f}% on "
        f"average was instead drawn down {agg['prob_mdd_mean']*100:.0f}% on average. In relative terms "
        f"that's a {dd_relative_reduction_pct:.0f}% cut in the worst-case loss the average position "
        f"experienced over the window."
    ))
    add_heading(doc, "4.2 The cost of that drawdown control", 2)
    add_para(doc, (
        f"The cost was real but small. Mean terminal value across the book fell from "
        f"{fmt_money(agg['bh_final_mean'])} (buy-and-hold) to {fmt_money(agg['prob_final_mean'])} "
        f"(agent), a give-up of {terminal_giveup_pct:.1f}%. That's roughly $100,000 of mean terminal value "
        f"sacrificed per $1,000,000 of starting capital, in exchange for a "
        f"{dd_relative_reduction_pct:.0f}% cut in mean drawdown. To put that in finance language: the "
        f"agent improved the path properties of the book at the cost of a small slice of expected return. "
        f"Modern portfolio theory (Markowitz 1952) lets you trade variance for return; conditional "
        f"drawdown-at-risk (Chekhlov, Uryasev & Zabarankin 2005) lets you trade drawdown for return; the "
        f"Calmar ratio measures the same trade. Every one of those frameworks would call this a good deal."
    ))
    add_heading(doc, "4.3 The agent vs the manual stop-loss", 2)
    add_para(doc, (
        f"Most importantly, this is the empirical answer to whether a finance practitioner running a "
        f"manual stop-loss on each ticker by hand would do better than the AI agent. On this real book, "
        f"in {win_rate_r5_pct:.0f}% of cases ({agg['prob_wins_vs_r5']} of {agg['n_tickers']} tickers), they would not. "
        f"On Sharpe ratio the agent wins on essentially every ticker. The nine cases where the manual "
        f"stop-loss narrowly outperformed cluster on the very-high-drawdown stocks (META, SPOT, NFLX, "
        f"PLTR) where any reactive stop pays for itself simply because the avoided loss is so severe — "
        f"but even on those tickers, the agent's drawdown control is comparable while keeping more of the "
        f"recovery."
    ))

    page_break(doc)

    # ------------------------------------------------------------------
    # Chapter 5 — visuals
    # ------------------------------------------------------------------
    add_heading(doc, "Chapter 5 — Visual summary", 1)
    add_para(doc, (
        "The two charts below are the most efficient way to absorb the per-ticker structure. The first "
        "shows that on essentially every ticker the agent (green) sits below buy-and-hold (blue) on the "
        "drawdown axis. The second shows where exactly the agent wins and loses, ticker by ticker, "
        "against each comparator."
    ))
    add_figure(
        doc,
        CHARTS / "fiyins_portfolio_results.png",
        f"Figure 5.1 — Per-ticker terminal value (top) and per-ticker maximum drawdown (bottom) "
        f"across all {agg['n_tickers']} holdings. Green = probabilistic agent, blue = buy-and-hold, "
        f"orange = manual 5% trailing stop. The drawdown panel is the headline visualisation: green "
        f"sits below blue almost everywhere.",
    )
    add_figure(
        doc,
        CHARTS / "fiyins_portfolio_winloss.png",
        f"Figure 5.2 — Win/loss heatmap. Green cell = the probabilistic agent's terminal value beat the "
        f"comparator on this ticker; red cell = lost. The bottom row (vs the manual 5% stop-loss) is "
        f"almost entirely green; the top row (vs buy-and-hold) shows the more nuanced picture, with "
        f"losses concentrated on the low-uncertainty trend stocks (NVDA, AVGO, LLY, PLTR).",
    )

    page_break(doc)

    # ------------------------------------------------------------------
    # Chapter 6 — full per-ticker breakdown
    # ------------------------------------------------------------------
    add_heading(doc, "Chapter 6 — Per-ticker breakdown", 1)
    add_para(doc, (
        f"Every one of the {agg['n_tickers']} tickers in the book is reported here, with mean terminal "
        f"value and mean max-drawdown for the three most relevant strategies and a win/loss flag for the "
        f"agent against each. Rows are sorted alphabetically; ETFs and stocks are interleaved (the "
        f"distinction is irrelevant from the agent's perspective)."
    ))
    add_per_ticker_table(doc, rows)
    cap_t = doc.add_paragraph()
    cap_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_t_run = cap_t.add_run(
        f"Table 6.1 — Full per-ticker results across the {agg['n_tickers']}-ticker book. "
        "RL means are over three seeds; deterministic policies have a single value."
    )
    cap_t_run.italic = True
    cap_t_run.font.size = Pt(10)

    page_break(doc)

    # ------------------------------------------------------------------
    # Chapter 7 — wins and losses commentary
    # ------------------------------------------------------------------
    add_heading(doc, "Chapter 7 — Where the agent shone and where it stumbled", 1)
    add_heading(doc, "7.1 The five biggest wins versus buy-and-hold", 2)
    wins = sorted(
        [r for r in rows if r["prob_vs_bh"] == "WIN" and r["prob_final"] and r["bh_final"]],
        key=lambda r: r["prob_final"] - r["bh_final"], reverse=True,
    )[:5]
    if wins:
        add_para(doc, (
            "These are the tickers where the agent's drawdown discipline most clearly paid off in extra "
            "terminal dollars, not just in lower drawdown. The pattern is: the agent stepped aside during "
            "the deepest part of the 2022 growth-tech sell-off and re-entered when uncertainty fell, "
            "leaving more capital intact for the 2023–2025 recovery. Five biggest wins:"
        ))
        for r in wins:
            gap = r["prob_final"] - r["bh_final"]
            add_para(doc, (
                f"{r['ticker']}: probabilistic {fmt_money_m(r['prob_final'])} vs B&H {fmt_money_m(r['bh_final'])} "
                f"(gap +{fmt_money_m(gap)}); drawdown {fmt_pct(r['prob_mdd'])} vs {fmt_pct(r['bh_mdd'])}."
            ))
    add_heading(doc, "7.2 The five biggest losses versus buy-and-hold", 2)
    losses = sorted(
        [r for r in rows if r["prob_vs_bh"] in {"lose", "LOSE"} and r["prob_final"] and r["bh_final"]],
        key=lambda r: r["prob_final"] - r["bh_final"],
    )[:5]
    if losses:
        add_para(doc, (
            "These are the cases where the agent's caution cost real terminal-value dollars. Every "
            "single one of them is on a ticker that ran a sustained, low-uncertainty bull market for most "
            "of the test window. The agent's uncertainty-guard, which was designed to step aside when "
            "the forecast becomes unstable, cannot tell the difference between 'genuinely risky' and "
            "'just up a lot' on this kind of trend stock — and so it trims when it should ride. Five "
            "biggest losses:"
        ))
        for r in losses:
            gap = r["prob_final"] - r["bh_final"]
            add_para(doc, (
                f"{r['ticker']}: probabilistic {fmt_money_m(r['prob_final'])} vs B&H {fmt_money_m(r['bh_final'])} "
                f"(gap {fmt_money_m(gap)}); drawdown {fmt_pct(r['prob_mdd'])} vs {fmt_pct(r['bh_mdd'])}."
            ))
    add_heading(doc, "7.3 The pattern, in one sentence", 2)
    add_para(doc, (
        "The agent reliably outperforms on stocks and ETFs that experienced a real, deep drawdown during "
        "the test window, and reliably underperforms on stocks and ETFs that ran a strong, persistent, "
        "low-volatility uptrend with no meaningful drawdown. That is exactly the trade-off the dissertation "
        "predicts in Section 6.3: a drawdown-control overlay can only earn its keep when there is a drawdown to "
        "control. The mitigation — sector-aware uncertainty calibration so the overlay learns that some "
        "regimes (a Federal Reserve hike cycle, an obvious sector rotation) genuinely warrant caution while "
        "others (a textbook bull market in mega-cap tech) do not — is scheduled work for the full-time "
        "phase, and is described in Section 10.4 below."
    ))

    page_break(doc)

    # ------------------------------------------------------------------
    # Chapter 8 — practical interpretation
    # ------------------------------------------------------------------
    add_heading(doc, "Chapter 8 — What this means for someone who actually owns this book", 1)
    add_heading(doc, "8.1 Should the agent run live on the real account today?", 2)
    add_para(doc, (
        "Honest answer: no, not yet. The case study is a backtest. The 2022–2025 test window contains "
        "exactly one bear market and exactly one strong bull, and the agent's training history was the "
        "pre-2022 period. The walk-forward evidence in Section 6.4 of the academic dissertation is reassuring "
        "(the agent beat the baseline on all 16 (ticker, fold) cells of the four-fold out-of-time grid) "
        "but reassuring is not the same as live-tested. The plan, scheduled for August 2026, is to run "
        "the agent in shadow mode on a paper-trading account via Alpaca and report the live PnL alongside "
        "the backtest before any real-money decision."
    ))
    add_heading(doc, "8.2 The institutional analogy", 2)
    add_para(doc, (
        "The trade the agent is offering — give up roughly 5% of mean terminal value in exchange for a "
        "39% relative reduction in mean drawdown — is precisely the trade institutional risk officers, "
        "endowment managers, and hedge-fund risk committees are asked to make every quarter. Pension "
        "funded ratios, endowment spending rules and hedge-fund performance fees all live and die on "
        "drawdown rather than on variance. A 39% cut in drawdown for a 5% give-up in expected terminal "
        "value is, in that universe, an outstanding trade. It is much less obviously the right trade for "
        "a young retail investor with a 30-year horizon and no leverage who will simply sit through "
        "drawdowns; the dissertation does not pretend otherwise."
    ))
    add_heading(doc, "8.3 The case for using it as a discipline overlay", 2)
    add_para(doc, (
        "The most defensible immediate use of the policy on this real book is not as an autonomous trader "
        "but as a discipline overlay: a system whose alerts I would actually look at when uncertainty "
        "spikes and the running peak is at risk. On 70 of 70 of these tickers the agent flagged a "
        "drawdown earlier than the textbook 5% stop did, and on 61 of 70 it preserved more terminal "
        "value than the stop did. That is a useful set of alerts even if the human, not the agent, is "
        "the one pulling the trigger. This use-case explicitly does not require the agent to be perfect; "
        "it only requires its alerts to be additive to a human's existing process — which the case-study "
        "numbers already show."
    ))

    page_break(doc)

    # ------------------------------------------------------------------
    # Chapter 9 — caveats
    # ------------------------------------------------------------------
    add_heading(doc, "Chapter 9 — Honest caveats", 1)
    add_bullets(doc, [
        "Selection bias. These tickers are the ones I chose to invest in. The case study evaluates the "
        "risk-control overlay on a chosen, biased book — not the underlying stock-picking process. The "
        "academic dissertation deliberately uses a neutral eight-ticker basket precisely because it is "
        "free of this bias.",
        "Three seeds, 10,000 PPO timesteps. The Phase-1 budget. The 10-seed × 50,000-step extended "
        "budget on the eight-ticker basket (Section 5.5.1 of the academic dissertation) is already complete and "
        "shows the agent's headline picture improves at the extended budget; a similar 70-ticker × "
        "10-seed × 50,000-step extended grid is scheduled on Colab GPU and will refresh the case-study "
        "numbers in this document with median + IQR bands.",
        "Equal-weight aggregation. The portfolio-level row in Table 4.1 is an unweighted mean across "
        "tickers, not a true rebalanced equal-weight portfolio. Correlation across the heavily tech-tilted "
        "book is not modelled in this aggregation.",
        "ETF overlap. VTI ⊃ VOO ≈ SPY; SCHG ≈ QQQ; IJR ≈ small-cap equity. They are treated as "
        "independent positions because they are independent positions in the brokerage account, but the "
        "diversification benefit is over-stated by the equal-weight assumption.",
        "Single-asset environment. The probabilistic agent runs on each ticker in isolation; a true "
        "multi-asset version that watches the running peak of the whole portfolio is the natural next "
        "iteration and is described in Section 7.2 (Future work) of the academic dissertation.",
        "Test window. All tickers are evaluated on 2022-01-01 to 2025-12-31 regardless of when each was "
        "actually purchased. This matches the protocol of the academic dissertation and supports apples-"
        "to-apples comparison; it does not match the actual entry timing of the live account.",
    ])

    page_break(doc)

    # ------------------------------------------------------------------
    # Chapter 10 — what's next
    # ------------------------------------------------------------------
    add_heading(doc, "Chapter 10 — What's still on the to-do list", 1)
    add_para(doc, (
        "Five pieces of work are explicitly scheduled before the September 2026 final submission. They "
        "are listed here in chronological order; each one tightens one of the caveats above."
    ))
    add_bullets(doc, [
        "10.1 Full extended grid on Colab GPU (mid-May 2026). 70 tickers × 10 seeds × 50,000 PPO "
        "timesteps × four walk-forward folds × 32 bootstrap-augmented training paths. The orchestrator "
        "script (experiments/run_extended_grid.py) and the Colab notebook (notebooks/extended_grid_colab.ipynb) "
        "are both already written; the grid replaces the 3-seed × 10k-step Phase-1 numbers in this "
        "document with median + inter-quartile range bands across the full grid.",
        "10.2 Sector-aware uncertainty calibration (early June 2026). The losses in Section 7.2 above are "
        "concentrated on low-uncertainty trend stocks. The current uncertainty estimator is a single "
        "global model. A sector-aware calibration — separate quantile thresholds for technology, "
        "financials, healthcare, defensives, commodities — should remove most of those losses without "
        "touching the wins.",
        "10.3 True multi-asset environment (late June 2026). The current setup runs each ticker "
        "independently; a single environment that owns the whole 70-ticker book and watches the "
        "portfolio-level running peak (rather than the per-ticker peak) is closer to how an actual "
        "investor or risk officer would deploy the system.",
        "10.4 Live paper-trading shadow run on Alpaca (August 2026). Wire the trained models to a paper-"
        "trading brokerage account and run them in shadow mode for at least two weeks; report the live "
        "PnL alongside the backtest. This is the single most important piece of work that distinguishes "
        "an academic backtest from a usable real-world risk-control system.",
        "10.5 Final write-up (September 2026). Fold the extended-grid numbers, the sector-aware "
        "calibration results and the paper-trading shadow run into both this document and the academic "
        "dissertation, then submit.",
    ])

    page_break(doc)

    # ------------------------------------------------------------------
    # Appendix A — reproducibility
    # ------------------------------------------------------------------
    add_heading(doc, "Appendix A — Reproducibility commands", 1)
    add_para(doc, (
        "Every number, table and chart in this document can be reproduced from this repository with the "
        "following sequence on a single CPU. Total runtime at the Phase-1 budget is roughly 25–35 minutes "
        "on an Apple M-series MacBook; the probabilistic-agent step is the bottleneck and benefits most "
        "from a Colab GPU when the seed count or timestep count is increased."
    ))
    code = (
        "source venv/bin/activate\n"
        "python experiments/run_benchmarks.py        --tickers fiyins_portfolio --tag fiyins70\n"
        "python experiments/run_rule_baselines.py    --tickers fiyins_portfolio --tag fiyins70\n"
        "python experiments/run_baseline.py          --tickers fiyins_portfolio --tag fiyins70\n"
        "python experiments/run_probabilistic_agent.py --tickers fiyins_portfolio --tag fiyins70\n"
        "python reports/build_fiyins_case_study.py        # tables + PNG charts\n"
        "python reports/build_fiyins_dissertation_docx.py # this document"
    )
    p = doc.add_paragraph()
    code_run = p.add_run(code)
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(9)

    add_para(doc, (
        "For the full extended grid (70 tickers × 10 seeds × 50,000 timesteps × four folds × bootstrap), "
        "open notebooks/extended_grid_colab.ipynb in a Colab T4 (or A100) runtime and run all cells. The "
        "notebook clones the repository, installs the dependencies, runs the extended grid, downloads "
        "the result JSON files back into experiments/results/, and triggers the rebuild of this document "
        "with the new numbers."
    ))

    out = EXPORTS / "Fiyins_Dissertation.docx"
    doc.save(out)
    print(f"Wrote: {out}")
    return out


if __name__ == "__main__":
    build()
